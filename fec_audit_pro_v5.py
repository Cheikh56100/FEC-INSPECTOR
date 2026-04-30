#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
FEC Audit Pro v5.0
Nouveautes v5 :
  - MODULE 1 : Generateur de Rapport de Mission IA (Word/PDF)
      Envoie les resultats d'analyse a l'API Claude pour rediger
      un rapport professionnel 3 pages destine au client final.
      Necessite : pip install python-docx anthropic
  - MODULE 2 : Bouclier Fiscal — Score de Probabilite de Controle DGFiP
      Combine Benford, fin d'exercice, aberrants, TVA et benchmark
      pour produire un score 0-100 avec grille de risques detaillee.
  - MODULE 3 : Convertisseur FEC en Previsionnel N+1
      Projette les donnees historiques (CA, charges, personnel)
      pour generer automatiquement un budget/previsionnel N+1.
      Export Word (.docx) inclus.
  - Toutes les fonctionnalites v4 conservees (15 controles)
"""

import pandas as pd
import numpy as np
import json, sys, os, math, re
from datetime import datetime
from collections import defaultdict, Counter
import argparse
import warnings
warnings.filterwarnings('ignore')

# Imports optionnels (v5 — modules IA et Word)
try:
    import anthropic as _anthropic_lib
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False

try:
    from docx import Document as _DocxDocument
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ══════════════════════════════════════════════════════════════
# BASE DE BENCHMARKS SECTORIELS INTEGREE
# Sources : Banque de France (ratios sectoriels 2022-2023),
#           INSEE (comptes nationaux), DGFiP (statistiques fiscales)
# Cles NAF estimees → ratios medians secteur
# ══════════════════════════════════════════════════════════════

BENCHMARKS = {
    # (code_secteur, label_secteur) : {
    #   'tm_med'  : taux marge brute median %
    #   'tm_low'  : quartile bas %          tm_high : quartile haut %
    #   'tr_med'  : taux resultat median %
    #   'cp_ca'   : charges personnel / CA median %
    #   'tx_tva'  : taux TVA courant attendu %
    #   'note'    : commentaire sectoriel
    # }
    'RESTAURATION': {
        'label': 'Restauration / Hotellerie (56xx)',
        'tm_med': 68.0, 'tm_low': 58.0, 'tm_high': 76.0,
        'tr_med': 4.5,  'cp_ca': 35.0,  'tx_tva': 10.0,
        'note': "TVA reduite 10% sur les repas servis. Marge brute elevee (pas d'achats revendus massifs). "
                "Charges personnel typiquement 30-40% du CA. Saisonnalite forte pour l'hotellerie."
    },
    'COMMERCE_DETAIL': {
        'label': 'Commerce de detail (47xx)',
        'tm_med': 32.0, 'tm_low': 22.0, 'tm_high': 42.0,
        'tr_med': 3.2,  'cp_ca': 18.0,  'tx_tva': 20.0,
        'note': "Marge brute plus faible : achats de marchandises importants. "
                "Rotation des stocks cle. TVA 20% standard sur la plupart des produits."
    },
    'COMMERCE_GROS': {
        'label': 'Commerce de gros / Negoce (46xx)',
        'tm_med': 18.0, 'tm_low': 10.0, 'tm_high': 28.0,
        'tr_med': 2.8,  'cp_ca': 10.0,  'tx_tva': 20.0,
        'note': "Marges tres faibles caracteristiques du negoce. Volume eleve compense la marge unitaire. "
                "Attention aux credits clients etendus (impact tresorerie)."
    },
    'BTP': {
        'label': 'BTP / Construction (41xx-43xx)',
        'tm_med': 22.0, 'tm_low': 14.0, 'tm_high': 32.0,
        'tr_med': 3.5,  'cp_ca': 28.0,  'tx_tva': 10.0,
        'note': "TVA reduite 10% sur travaux de renovation logement. Marges variables selon sous-traitance. "
                "Suivi des chantiers en cours determinant pour le resultat."
    },
    'SERVICES_B2B': {
        'label': 'Services aux entreprises / Conseil (69xx-74xx)',
        'tm_med': 85.0, 'tm_low': 75.0, 'tm_high': 95.0,
        'tr_med': 12.0, 'cp_ca': 55.0,  'tx_tva': 20.0,
        'note': "Activite de service pure : quasiment pas d'achats revendus, marge brute proche de 100%. "
                "Les charges de personnel sont le poste dominant. Valoriser les encours clients."
    },
    'PROFESSION_LIBERALE': {
        'label': 'Professions liberales / Sante (86xx-69xx)',
        'tm_med': 90.0, 'tm_low': 82.0, 'tm_high': 97.0,
        'tr_med': 20.0, 'cp_ca': 30.0,  'tx_tva': 0.0,
        'note': "Activite souvent exoneree de TVA (medecins, avocats, experts-comptables). "
                "Marge brute proche de 100%. Charges de personnel moderees si exercice individuel."
    },
    'INDUSTRIE': {
        'label': 'Industrie manufacturiere (10xx-33xx)',
        'tm_med': 28.0, 'tm_low': 18.0, 'tm_high': 40.0,
        'tr_med': 4.8,  'cp_ca': 22.0,  'tx_tva': 20.0,
        'note': "Marge brute impactee par les matieres premieres et la sous-traitance. "
                "Amortissements importants (machines, outillage). Surveiller le BFR."
    },
    'TRANSPORT': {
        'label': 'Transport / Logistique (49xx)',
        'tm_med': 35.0, 'tm_low': 25.0, 'tm_high': 48.0,
        'tr_med': 3.0,  'cp_ca': 32.0,  'tx_tva': 20.0,
        'note': "Carburant et entretien vehicules = postes cles en 60x. "
                "Amortissements vehicules significatifs. Marge sous pression concurrentielle."
    },
    'IMMOBILIER': {
        'label': 'Immobilier / Agence (68xx)',
        'tm_med': 75.0, 'tm_low': 60.0, 'tm_high': 88.0,
        'tr_med': 8.0,  'cp_ca': 42.0,  'tx_tva': 20.0,
        'note': "CA constitue de commissions et honoraires. Saisonnalite marquee. "
                "Charges de personnel elevees (agents commerciaux)."
    },
    'INFORMATIQUE': {
        'label': 'Informatique / ESN (62xx)',
        'tm_med': 80.0, 'tm_low': 68.0, 'tm_high': 92.0,
        'tr_med': 8.5,  'cp_ca': 60.0,  'tx_tva': 20.0,
        'note': "Activite de service : peu d'achats revendus. "
                "Personnel hautement qualifie = masse salariale dominante. "
                "Attention aux provisions pour conges et RTT."
    },
    'INCONNU': {
        'label': 'Secteur non determine',
        'tm_med': 45.0, 'tm_low': 20.0, 'tm_high': 85.0,
        'tr_med': 5.0,  'cp_ca': 30.0,  'tx_tva': 20.0,
        'note': "Secteur non identifie automatiquement. Les benchmarks affiches sont des moyennes "
                "tous secteurs confondus (source Banque de France). Preciser le secteur pour une "
                "comparaison pertinente."
    },
}

# Liste pour le dialogue interactif
SECTEURS_LISTE = [
    ('1', 'RESTAURATION',      'Restauration / Hotellerie / Traiteur'),
    ('2', 'COMMERCE_DETAIL',   'Commerce de detail (magasin, e-commerce)'),
    ('3', 'COMMERCE_GROS',     'Commerce de gros / Negoce / Distribution'),
    ('4', 'BTP',               'BTP / Construction / Renovation'),
    ('5', 'SERVICES_B2B',      'Conseil / Services aux entreprises / Cabinet'),
    ('6', 'PROFESSION_LIBERALE','Profession liberale / Sante / Juridique'),
    ('7', 'INDUSTRIE',         'Industrie / Fabrication / Transformation'),
    ('8', 'TRANSPORT',         'Transport / Logistique / Livraison'),
    ('9', 'IMMOBILIER',        'Immobilier / Agence immobiliere'),
    ('10','INFORMATIQUE',      'Informatique / ESN / Developpement / SaaS'),
    ('0', 'INCONNU',           'Autre / Ne pas comparer au marche'),
]

# ══════════════════════════════════════════════════════════════
# DETECTION AUTOMATIQUE DU SECTEUR
# ══════════════════════════════════════════════════════════════

def detecter_secteur(df):
    """
    Analyse les libelles de comptes et les numeros de compte pour
    estimer le secteur d'activite. Retourne (code_secteur, score_confiance).
    """
    scores = {k: 0 for k in BENCHMARKS if k != 'INCONNU'}

    # — Collecte des libelles —
    libs = ' '.join(
        df['EcritureLib'].fillna('').tolist() +
        df['CompteLib'].fillna('').tolist()
    ).upper()

    def s(pfx, col='Debit'):
        return df[df['CompteNum'].str.startswith(pfx)][col].sum()

    ca_total  = s('70', 'Credit') - s('70', 'Debit')
    ach_total = s('60', 'Debit')
    cp_total  = s('64', 'Debit')
    tm_brute  = (ca_total - ach_total) / ca_total * 100 if ca_total > 0 else 50

    # — Mots-cles dans les libelles —
    kw_resto   = ['RESTAURANT','REPAS','TRAITEUR','HOTEL','HEBERGEMENT','BRASSERIE',
                  'CUISINE','LIVRAISON REPAS','UBER EATS','DELIVEROO','MENU']
    kw_btp     = ['CHANTIER','TRAVAUX','BATIMENT','CONSTRUCTION','RENOVATION','MACONNERIE',
                  'ELECTRICITE','PLOMBERIE','PEINTURE','COUVERTURE','GROS OEUVRE']
    kw_conseil = ['HONORAIRES','CONSEIL','CONSULTING','MISSION','PRESTATION','FORMATION',
                  'AUDIT','EXPERTISE','ASSISTANCE','MANAGEMENT']
    kw_info    = ['LICENCE','ABONNEMENT LOGICIEL','HEBERGEMENT WEB','MAINTENANCE INFORMATIQUE',
                  'DEVELOPPEMENT','SAAS','CLOUD','SERVEUR','INFRA']
    kw_sante   = ['CABINET MEDICAL','CONSULTATION','ORDONNANCE','PHARMACIE','CLINIQUE',
                  'SOIN','KINESITHERAPIE','DENTISTE','MEDECIN']
    kw_immo    = ['COMMISSION AGENCE','MANDAT','TRANSACTION IMMO','LOCATION','BAIL',
                  'LOYER','CHARGES LOCATIVES']
    kw_transport=['CARBURANT','GASOIL','PEAGE','VEHICULE','FLOTTE','TRANSPORT','LIVRAISON',
                  'CHAUFFEUR','FRET','EXPEDITION']

    def hit(kws): return sum(1 for k in kws if k in libs)

    scores['RESTAURATION']       += hit(kw_resto) * 3
    scores['BTP']                 += hit(kw_btp) * 3
    scores['SERVICES_B2B']        += hit(kw_conseil) * 3
    scores['INFORMATIQUE']        += hit(kw_info) * 3
    scores['PROFESSION_LIBERALE'] += hit(kw_sante) * 3
    scores['IMMOBILIER']          += hit(kw_immo) * 3
    scores['TRANSPORT']           += hit(kw_transport) * 3

    # — Signaux comptables (numeros de comptes dominants) —
    cpts = df['CompteNum'].str[:2].value_counts()

    # Fortes achats marchandises (60) = commerce ou industrie
    if s('601', 'Debit') > s('604', 'Debit') * 2:
        scores['COMMERCE_DETAIL'] += 4
        scores['COMMERCE_GROS']   += 3
    if s('604', 'Debit') > 0:  # achats d'etudes et services
        scores['SERVICES_B2B'] += 2
        scores['INFORMATIQUE'] += 1

    # Taux de marge brute tres eleve → services
    if tm_brute > 80:
        scores['SERVICES_B2B']        += 4
        scores['INFORMATIQUE']        += 3
        scores['PROFESSION_LIBERALE'] += 3
        scores['IMMOBILIER']          += 2
    elif tm_brute > 55:
        scores['RESTAURATION']        += 3
        scores['TRANSPORT']           += 1
    elif tm_brute < 30:
        scores['COMMERCE_GROS']   += 4
        scores['INDUSTRIE']        += 3
        scores['COMMERCE_DETAIL'] += 2

    # Charges personnel / CA elevees → services intensifs en main d'oeuvre
    if ca_total > 0:
        ratio_cp = cp_total / ca_total * 100
        if ratio_cp > 50:
            scores['INFORMATIQUE']    += 3
            scores['SERVICES_B2B']    += 2
        if ratio_cp > 35:
            scores['RESTAURATION'] += 2

    # Comptes 86xx (sante) ou 69xx (juridique)
    if '86' in cpts.index or '69' in cpts.index:
        scores['PROFESSION_LIBERALE'] += 5

    # Comptes 604 (achats prestations) dominant
    prest = s('604', 'Debit')
    if prest > ach_total * 0.6 and ach_total > 0:
        scores['SERVICES_B2B'] += 3
        scores['BTP'] += 2

    # — Selection du meilleur score —
    best = max(scores, key=lambda k: scores[k])
    total_s = sum(scores.values())
    conf = round(scores[best] / total_s * 100) if total_s > 0 else 0

    return best, conf, scores

# ══════════════════════════════════════════════════════════════
# DIALOGUE INTERACTIF SELECTION SECTEUR
# ══════════════════════════════════════════════════════════════

def demander_secteur(detected_code, detected_conf):
    """
    Propose le secteur detecte et laisse l'utilisateur confirmer ou corriger.
    Mode non-interactif : retourne le secteur detecte.
    """
    if not sys.stdin.isatty():
        return detected_code  # mode non-interactif (pipe, script)

    print(f"\n{'─'*60}")
    print("  DETECTION DU SECTEUR D'ACTIVITE")
    print(f"{'─'*60}")
    label_det = BENCHMARKS.get(detected_code, {}).get('label', detected_code)
    print(f"\n  Secteur estime : {label_det}")
    print(f"  Confiance       : {detected_conf}%")
    print()
    print("  Confirmer ou choisir un autre secteur :")
    print()
    for num, code, label in SECTEURS_LISTE:
        marker = "►" if code == detected_code else " "
        print(f"  {marker} [{num:>2}] {label}")
    print()

    while True:
        try:
            choix = input("  Votre choix (Entree = confirmer detecte) : ").strip()
            if choix == '':
                print(f"\n  ✓ Secteur confirme : {label_det}\n")
                return detected_code
            for num, code, label in SECTEURS_LISTE:
                if choix == num:
                    print(f"\n  ✓ Secteur selectionne : {label}\n")
                    return code
            print("  Choix invalide, recommencez.")
        except (EOFError, KeyboardInterrupt):
            return detected_code

# ══════════════════════════════════════════════════════════════
# CHARGEMENT FEC
# ══════════════════════════════════════════════════════════════

def detecter_sep(path):
    with open(path, 'r', encoding='utf-8-sig', errors='replace') as f:
        l = f.readline()
    for s in ['\t', '|', ';', ',']:
        if s in l: return s
    return '\t'

def charger_fec(path):
    sep = detecter_sep(path)
    df = None
    for enc in ['utf-8-sig', 'latin-1', 'cp1252']:
        try:
            df = pd.read_csv(path, sep=sep, encoding=enc, dtype=str,
                             on_bad_lines='skip', low_memory=False)
            break
        except:
            continue
    if df is None: raise ValueError("Impossible de lire le fichier FEC")
    df.columns = [c.strip().replace('\ufeff', '') for c in df.columns]
    for col in ['Debit', 'Credit', 'Montantdevise']:
        if col in df.columns:
            df[col] = df[col].fillna('0').str.replace(',', '.').str.replace(' ', '')
            df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)
    for col in ['EcritureDate', 'PieceDate', 'DateLet', 'ValidDate']:
        if col in df.columns:
            df[col] = pd.to_datetime(df[col], format='%Y%m%d', errors='coerce')
    for col in df.select_dtypes(include='object').columns:
        df[col] = df[col].fillna('').astype(str).str.strip()
    if 'CompteNum' not in df.columns:
        raise ValueError("Colonne CompteNum manquante — verifiez le format FEC")
    df['montant'] = (df.get('Debit', pd.Series(0, index=df.index)) +
                     df.get('Credit', pd.Series(0, index=df.index)))
    return df

# ══════════════════════════════════════════════════════════════
# CONTROLES EXISTANTS (identiques v3)
# ══════════════════════════════════════════════════════════════

def analyser_equilibre(df):
    anomalies = []
    if 'EcritureNum' not in df.columns: return anomalies, {}
    eq = df.groupby('EcritureNum').agg(d=('Debit', 'sum'), c=('Credit', 'sum'))
    eq['ecart'] = abs(eq['d'] - eq['c'])
    deseq = eq[eq['ecart'] > 0.02]
    for num, row in deseq.head(20).iterrows():
        anomalies.append({'type': 'Desequilibre D/C', 'gravite': 'CRITIQUE',
            'detail': f"Ecriture {num} : D={row['d']:,.2f} EUR / C={row['c']:,.2f} EUR (ecart {row['ecart']:,.2f} EUR)",
            'montant': float(row['ecart']), 'reference': str(num)})
    return anomalies, {'nb': len(deseq), 'ecart_total': float(deseq['ecart'].sum())}

def analyser_benford(df):
    attendus = {str(d): round(math.log10(1 + 1/d) * 100, 2) for d in range(1, 10)}
    m = df[df['montant'] > 0]['montant']
    prem = m.astype(str).str.replace(r'^0\.0*', '', regex=True).str[0]
    prem = prem[prem.str.match(r'[1-9]')]
    total = len(prem)
    if total < 100:
        return [], {'erreur': 'Moins de 100 montants, Benford non applicable', 'total': total}
    obs = {str(d): round((prem == str(d)).sum() / total * 100, 2) for d in range(1, 10)}
    chi2 = sum((obs[d] - attendus[d])**2 / attendus[d] for d in obs if attendus[d] > 0)
    anomalies = []
    if chi2 > 15.5:
        anomalies.append({'type': 'Loi de Benford violee', 'gravite': 'ALERTE',
            'detail': f"Chi2={chi2:.1f} (seuil 15.5) — distribution anormale, possible manipulation",
            'montant': 0, 'reference': 'Benford'})
    for d in [str(i) for i in range(1, 10)]:
        e = abs(obs[d] - attendus[d])
        if e > 5:
            anomalies.append({'type': f'Benford chiffre {d} suspect', 'gravite': 'ATTENTION',
                'detail': f"Chiffre {d} : observe {obs[d]:.1f}% vs attendu {attendus[d]:.1f}% (ecart {e:.1f} pts)",
                'montant': 0, 'reference': f'Chiffre {d}'})
    return anomalies, {'obs': obs, 'att': attendus, 'chi2': round(chi2, 2), 'total': total}

def analyser_montants_ronds(df):
    anomalies = []
    m = df[df['montant'] > 100]['montant']
    if len(m) == 0: return anomalies, {}
    total = len(m)
    r1000 = (m % 1000 == 0).sum()
    r100  = (m % 100 == 0).sum()
    p1000 = r1000 / total * 100
    p100  = r100 / total * 100
    if p1000 > 20:
        anomalies.append({'type': 'Montants ronds excessifs', 'gravite': 'ALERTE',
            'detail': f"{r1000} montants multiples de 1000 EUR ({p1000:.1f}%) — estimations ou ecritures fictives",
            'montant': float(m[m % 1000 == 0].sum()), 'reference': 'Montants'})
    elif p100 > 40:
        anomalies.append({'type': 'Montants ronds frequents', 'gravite': 'ATTENTION',
            'detail': f"{r100} montants multiples de 100 EUR ({p100:.1f}%) — a verifier",
            'montant': float(m[m % 100 == 0].sum()), 'reference': 'Montants'})
    top = df[df['montant'] % 1000 == 0].nlargest(5, 'montant')[
        ['CompteNum', 'CompteLib', 'EcritureLib', 'montant']].to_dict('records')
    return anomalies, {'p1000': round(p1000, 1), 'p100': round(p100, 1), 'top': top}

def analyser_weekend(df):
    anomalies = []
    if 'EcritureDate' not in df.columns: return anomalies, {}
    dd = df[df['EcritureDate'].notna()].copy()
    dd['dow'] = dd['EcritureDate'].dt.dayofweek
    we = dd[dd['dow'].isin([5, 6])]
    nb = len(we); total = len(dd)
    pct = nb / total * 100 if total > 0 else 0
    par_jour = dd.groupby('dow').size().to_dict()
    jours = ['Lundi', 'Mardi', 'Mercredi', 'Jeudi', 'Vendredi', 'Samedi', 'Dimanche']
    if nb > 0 and pct > 5:
        anomalies.append({'type': 'Ecritures week-end', 'gravite': 'ATTENTION',
            'detail': f"{nb} ecriture(s) le week-end ({pct:.1f}% du total) — montant {float(we['montant'].sum()):,.0f} EUR",
            'montant': float(we['montant'].sum()), 'reference': 'Dates'})
    return anomalies, {'nb': nb, 'pct': round(pct, 1),
                       'par_jour': {jours[j]: int(n) for j, n in par_jour.items()}}

def analyser_soldes(df):
    anomalies = []
    sol = df.groupby(['CompteNum', 'CompteLib']).agg(
        d=('Debit', 'sum'), c=('Credit', 'sum'), nb=('Debit', 'count')).reset_index()
    sol['solde'] = sol['d'] - sol['c']
    sol['cl']  = sol['CompteNum'].str[:1]
    sol['scl'] = sol['CompteNum'].str[:3]
    for _, r in sol[(sol['cl'] == '6') & (sol['solde'] < -100)].iterrows():
        anomalies.append({'type': 'Charge creditrice', 'gravite': 'ALERTE',
            'detail': f"Cpt {r['CompteNum']} ({r['CompteLib']}) : solde crediteur {abs(r['solde']):,.0f} EUR",
            'montant': abs(float(r['solde'])), 'reference': r['CompteNum']})
    for _, r in sol[(sol['cl'] == '7') & (sol['solde'] > 100)].iterrows():
        anomalies.append({'type': 'Produit debiteur', 'gravite': 'ALERTE',
            'detail': f"Cpt {r['CompteNum']} ({r['CompteLib']}) : solde debiteur {r['solde']:,.0f} EUR",
            'montant': float(r['solde']), 'reference': r['CompteNum']})
    for _, r in sol[(sol['scl'] == '401') & (sol['solde'] > 500)].iterrows():
        anomalies.append({'type': 'Fournisseur debiteur', 'gravite': 'ATTENTION',
            'detail': f"Cpt {r['CompteNum']} ({r['CompteLib']}) : solde debiteur {r['solde']:,.0f} EUR",
            'montant': float(r['solde']), 'reference': r['CompteNum']})
    for _, r in sol[(sol['scl'] == '411') & (sol['solde'] < -500)].iterrows():
        anomalies.append({'type': 'Client crediteur', 'gravite': 'ATTENTION',
            'detail': f"Cpt {r['CompteNum']} ({r['CompteLib']}) : solde crediteur {abs(r['solde']):,.0f} EUR",
            'montant': abs(float(r['solde'])), 'reference': r['CompteNum']})
    return anomalies, {'soldes': sol[['CompteNum', 'CompteLib', 'd', 'c', 'solde']].rename(
        columns={'d': 'debit', 'c': 'credit'}).to_dict('records')}

def analyser_doublons(df):
    anomalies = []
    cles = [c for c in ['EcritureDate', 'CompteNum', 'Debit', 'Credit'] if c in df.columns]
    if len(cles) < 3: return anomalies, {'nb': 0}
    dup = df[df.duplicated(subset=cles, keep=False)]
    if len(dup) > 0:
        mt = float(dup['montant'].sum())
        anomalies.append({'type': 'Doublons detectes', 'gravite': 'ALERTE',
            'detail': f"{len(dup)} lignes en doublon ({len(dup.groupby(cles))} groupes) — {mt:,.0f} EUR cumule",
            'montant': mt, 'reference': 'Multi'})
    return anomalies, {'nb': len(dup)}

def analyser_inversees(df):
    anomalies = []
    if 'CompteNum' not in df.columns: return anomalies, {'nb': 0}
    pos = df[df['Debit'] > 0][['CompteNum', 'EcritureDate', 'Debit']].copy()
    neg = df[df['Credit'] > 0][['CompteNum', 'EcritureDate', 'Credit']].copy()
    neg.columns = ['CompteNum', 'EcritureDate', 'Debit']
    m = pd.merge(pos, neg, on=['CompteNum', 'Debit'])
    if len(m) > 0:
        anomalies.append({'type': 'Ecritures inversees', 'gravite': 'ATTENTION',
            'detail': f"{len(m)} ecriture(s) avec montant identique en sens inverse — extournes possibles",
            'montant': float(m['Debit'].sum()), 'reference': 'Extourne'})
    return anomalies, {'nb': len(m)}

def analyser_fin_exercice(df):
    anomalies = []
    if 'EcritureDate' not in df.columns: return anomalies, {}
    dates = df['EcritureDate'].dropna()
    if len(dates) == 0: return anomalies, {}
    dmax = dates.max()
    d7  = df[df['EcritureDate'] >= dmax - pd.Timedelta(days=7)]
    d30 = df[df['EcritureDate'] >= dmax - pd.Timedelta(days=30)]
    total = float(df['montant'].sum())
    p7  = float(d7['montant'].sum())  / total * 100 if total > 0 else 0
    p30 = float(d30['montant'].sum()) / total * 100 if total > 0 else 0
    if p7 > 25:
        anomalies.append({'type': 'Concentration fin exercice', 'gravite': 'ALERTE',
            'detail': f"{p7:.1f}% du volume sur les 7 derniers jours — habillage de bilan possible",
            'montant': float(d7['montant'].sum()), 'reference': 'Cloture'})
    elif p30 > 50:
        anomalies.append({'type': 'Concentration fin de periode', 'gravite': 'ATTENTION',
            'detail': f"{p30:.1f}% du volume sur les 30 derniers jours — a controler",
            'montant': float(d30['montant'].sum()), 'reference': 'Cloture'})
    if 'EcritureNum' in df.columns:
        nums = pd.to_numeric(df['EcritureNum'], errors='coerce').dropna().astype(int)
        if len(nums) > 0:
            manq = set(range(nums.min(), nums.max() + 1)) - set(nums)
            if len(manq) > 0:
                anomalies.append({'type': 'Trous dans la numerotation', 'gravite': 'ALERTE',
                    'detail': f"{len(manq)} numero(s) manquant(s) entre {nums.min()} et {nums.max()} — suppressions possibles",
                    'montant': 0, 'reference': 'EcritureNum'})
    return anomalies, {'p7': round(p7, 1), 'p30': round(p30, 1),
        'date_cloture': str(dmax.date()) if dmax and not pd.isna(dmax) else 'N/A'}

def analyser_concentration(df):
    anomalies = []
    ca = df[df['CompteNum'].str.startswith('70')]['Credit'].sum()
    if ca == 0: return anomalies, {'top_clients': [], 'top_fournisseurs': []}
    tc = []
    if 'CompAuxNum' in df.columns:
        clients = df[df['CompteNum'].str.startswith('411')].copy()
        if clients['CompAuxNum'].str.len().max() > 0:
            gc = clients.groupby(['CompAuxNum', 'CompAuxLib'])['Credit'].sum().reset_index()
            gc['pct'] = gc['Credit'] / ca * 100
            gc = gc.sort_values('pct', ascending=False)
            tc = gc.head(5).to_dict('records')
            if len(gc) > 0 and gc.iloc[0]['pct'] > 50:
                r = gc.iloc[0]
                anomalies.append({'type': 'Client dominant', 'gravite': 'ALERTE',
                    'detail': f"Client {r['CompAuxNum']} ({r['CompAuxLib']}) : {r['pct']:.1f}% du CA — dependance elevee",
                    'montant': float(r['Credit']), 'reference': str(r['CompAuxNum'])})
    achats = df[df['CompteNum'].str.startswith('60')]['Debit'].sum()
    tf = []
    if 'CompAuxNum' in df.columns and achats > 0:
        four = df[df['CompteNum'].str.startswith('401')].copy()
        if four['CompAuxNum'].str.len().max() > 0:
            gf = four.groupby(['CompAuxNum', 'CompAuxLib'])['Debit'].sum().reset_index()
            gf['pct'] = gf['Debit'] / achats * 100
            gf = gf.sort_values('pct', ascending=False)
            tf = gf.head(5).to_dict('records')
            if len(gf) > 0 and gf.iloc[0]['pct'] > 60:
                r = gf.iloc[0]
                anomalies.append({'type': 'Fournisseur dominant', 'gravite': 'ATTENTION',
                    'detail': f"Fournisseur {r['CompAuxNum']} : {r['pct']:.1f}% des achats",
                    'montant': float(r['Debit']), 'reference': str(r['CompAuxNum'])})
    return anomalies, {'top_clients': tc, 'top_fournisseurs': tf}

def analyser_saisonnalite(df):
    anomalies = []
    if 'EcritureDate' not in df.columns: return anomalies, {}
    dd = df[df['EcritureDate'].notna() & df['CompteNum'].str.startswith('70')].copy()
    if len(dd) == 0: return anomalies, {}
    dd['mois'] = dd['EcritureDate'].dt.month
    cam = dd.groupby('mois')['Credit'].sum()
    cv = cam.std() / cam.mean() * 100 if len(cam) >= 6 and cam.mean() > 0 else 999
    if cv < 10 and len(cam) >= 6:
        anomalies.append({'type': 'CA trop regulier', 'gravite': 'ATTENTION',
            'detail': f"CV mensuel : {cv:.1f}% — CA quasi-constant, possible lissage artificiel",
            'montant': 0, 'reference': 'CA mensuel'})
    vv = df[df['EcritureDate'].notna()].copy()
    vv['periode'] = vv['EcritureDate'].dt.to_period('M')
    vol = vv.groupby('periode').agg(nb=('montant', 'count'), mt=('montant', 'sum')).reset_index()
    return anomalies, {
        'ca_mensuel':  {str(k): round(float(v), 2) for k, v in cam.items()},
        'vol_mensuel': [(str(r['periode']), int(r['nb']), round(float(r['mt']), 2))
                        for _, r in vol.iterrows()]
    }

def analyser_marges(df):
    anomalies = []
    def s(pfx, col):
        return df[df['CompteNum'].str.startswith(pfx)][col].sum()
    ca     = s('70', 'Credit') - s('70', 'Debit')
    achats = s('60', 'Debit')  - s('60', 'Credit')
    cp     = s('64', 'Debit')
    ce     = s('61', 'Debit') + s('62', 'Debit')
    tp     = s('7', 'Credit') - s('7', 'Debit')
    tc     = s('6', 'Debit')  - s('6', 'Credit')
    mb     = ca - achats
    tm     = mb / ca * 100 if ca > 0 else 0
    res    = tp - tc
    tr     = res / tp * 100 if tp > 0 else 0
    tva_c  = s('44571', 'Credit')
    tva_d  = df[df['CompteNum'].str.startswith(('44566', '44562', '44563'))]['Debit'].sum()
    tx_tva = tva_c / ca * 100 if ca > 0 else 0
    if ca > 0 and tm < 0:
        anomalies.append({'type': 'Marge negative', 'gravite': 'CRITIQUE',
            'detail': f"Taux de marge brute : {tm:.1f}% — ventes inferieures aux achats",
            'montant': abs(float(mb)), 'reference': '70x/60x'})
    if ca > 0 and cp > ca:
        anomalies.append({'type': 'Charges personnel > CA', 'gravite': 'CRITIQUE',
            'detail': f"Charges personnel ({cp:,.0f} EUR) superieures au CA ({ca:,.0f} EUR)",
            'montant': float(cp - ca), 'reference': '64x'})
    if ca > 0 and tx_tva > 25:
        anomalies.append({'type': 'TVA incoherente', 'gravite': 'ALERTE',
            'detail': f"Taux TVA apparent {tx_tva:.1f}% superieur au maximum legal (20%)",
            'montant': float(tva_c), 'reference': '4457x'})
    if tva_d > tva_c * 1.5 and tva_c > 0:
        anomalies.append({'type': 'TVA deductible excessive', 'gravite': 'ALERTE',
            'detail': f"TVA deductible ({tva_d:,.0f}) >> TVA collectee ({tva_c:,.0f}) — credit de TVA a verifier",
            'montant': float(tva_d - tva_c), 'reference': '4456x'})
    return anomalies, {
        'ca': float(ca), 'achats': float(achats), 'mb': float(mb), 'tm': float(tm),
        'cp': float(cp), 'ce': float(ce), 'tp': float(tp), 'tc': float(tc),
        'res': float(res), 'tr': float(tr),
        'tva_c': float(tva_c), 'tva_d': float(tva_d), 'tx_tva': float(tx_tva)
    }

def analyser_aberrants(df):
    anomalies = []
    m = df[df['montant'] > 0]['montant']
    if len(m) < 10: return anomalies, {}
    mean, std = float(m.mean()), float(m.std())
    if std == 0: return anomalies, {}
    df2 = df.copy()
    df2['zs'] = (df2['montant'] - mean) / std
    ab = df2[(df2['zs'].abs() > 3) & (df2['montant'] > 5000)].nlargest(8, 'zs')
    for _, r in ab.iterrows():
        anomalies.append({'type': 'Montant statistiquement aberrant', 'gravite': 'ATTENTION',
            'detail': f"{r['montant']:,.0f} EUR — z={r['zs']:.1f} ecarts-types — {r.get('CompteLib', '')} / {r.get('EcritureLib', '')}",
            'montant': float(r['montant']), 'reference': r.get('CompteNum', '')})
    return anomalies, {'mean': round(mean, 2), 'std': round(std, 2), 'nb': len(ab)}

# ══════════════════════════════════════════════════════════════
# NOUVEAUX MODULES v4
# ══════════════════════════════════════════════════════════════

def analyser_benchmark(df, secteur_code):
    """
    Compare les ratios de l'entreprise avec les benchmarks sectoriels.
    Genere des anomalies lorsque les ecarts sont significatifs.
    """
    anomalies = []
    bench = BENCHMARKS.get(secteur_code, BENCHMARKS['INCONNU'])

    def s(pfx, col):
        return df[df['CompteNum'].str.startswith(pfx)][col].sum()

    ca     = s('70', 'Credit') - s('70', 'Debit')
    achats = s('60', 'Debit')  - s('60', 'Credit')
    cp     = s('64', 'Debit')
    mb     = ca - achats
    tm     = mb / ca * 100 if ca > 0 else None
    ratio_cp = cp / ca * 100 if ca > 0 else None
    tva_c  = s('44571', 'Credit')
    tx_tva = tva_c / ca * 100 if ca > 0 else 0

    resultats = {
        'secteur_code':  secteur_code,
        'secteur_label': bench['label'],
        'tm_entreprise': round(tm, 1) if tm is not None else None,
        'tm_mediane':    bench['tm_med'],
        'tm_low':        bench['tm_low'],
        'tm_high':       bench['tm_high'],
        'cp_ca_entreprise': round(ratio_cp, 1) if ratio_cp is not None else None,
        'cp_ca_mediane':    bench['cp_ca'],
        'tx_tva_attendu':   bench['tx_tva'],
        'tx_tva_observe':   round(tx_tva, 1),
        'note_sectorielle': bench['note'],
        'ca': float(ca),
    }

    if ca <= 0:
        return anomalies, resultats

    # — Comparaison taux de marge brute —
    if tm is not None:
        ecart_tm = tm - bench['tm_med']
        if tm < bench['tm_low']:
            anomalies.append({'type': 'Marge sous la mediane sectorielle', 'gravite': 'ALERTE',
                'detail': (f"Marge brute : {tm:.1f}% vs mediane secteur {bench['tm_med']:.0f}% "
                           f"[{bench['tm_low']:.0f}%-{bench['tm_high']:.0f}%] — "
                           f"ecart de {abs(ecart_tm):.1f} pts sous le bas de fourchette. "
                           f"Risque : sous-evaluation des recettes ou sur-evaluation des achats."),
                'montant': 0, 'reference': 'Benchmark'})
        elif tm > bench['tm_high']:
            anomalies.append({'type': 'Marge au-dessus de la mediane sectorielle', 'gravite': 'ATTENTION',
                'detail': (f"Marge brute : {tm:.1f}% vs mediane secteur {bench['tm_med']:.0f}% "
                           f"[{bench['tm_low']:.0f}%-{bench['tm_high']:.0f}%] — "
                           f"ecart de {abs(ecart_tm):.1f} pts au-dessus du haut de fourchette. "
                           f"A verifier : achats sous-estimes ou activite atypique."),
                'montant': 0, 'reference': 'Benchmark'})

    # — Comparaison charges de personnel —
    if ratio_cp is not None:
        ecart_cp = ratio_cp - bench['cp_ca']
        if ecart_cp > 15:
            anomalies.append({'type': 'Charges personnel anormalement elevees', 'gravite': 'ATTENTION',
                'detail': (f"Charges personnel/CA : {ratio_cp:.1f}% vs mediane secteur {bench['cp_ca']:.0f}% "
                           f"— ecart +{ecart_cp:.1f} pts. Verifier si des salaires fictifs, "
                           f"des dirigeants sur-remuneres ou des charges de personnel hors activite."),
                'montant': float(cp), 'reference': '64x'})
        elif ecart_cp < -20 and cp > 0:
            anomalies.append({'type': 'Charges personnel anormalement faibles', 'gravite': 'INFO',
                'detail': (f"Charges personnel/CA : {ratio_cp:.1f}% vs mediane secteur {bench['cp_ca']:.0f}% "
                           f"— ecart {ecart_cp:.1f} pts. Possible sous-declaration, "
                           f"recours massif a la sous-traitance ou activite saisonniere."),
                'montant': 0, 'reference': '64x'})

    # — Coherence TVA —
    if bench['tx_tva'] == 0 and tx_tva > 1:
        anomalies.append({'type': 'TVA inattendue (secteur exonere)', 'gravite': 'ALERTE',
            'detail': (f"Le secteur {bench['label']} est habituellement exonere de TVA "
                       f"mais un taux TVA apparent de {tx_tva:.1f}% est detecte. "
                       f"Verifier le regime TVA de la societe."),
            'montant': float(tva_c), 'reference': '4457x'})
    elif bench['tx_tva'] > 0:
        ecart_tva = abs(tx_tva - bench['tx_tva'])
        if tx_tva > 0 and ecart_tva > 8:
            anomalies.append({'type': 'Taux TVA incoherent avec le secteur', 'gravite': 'ATTENTION',
                'detail': (f"TVA apparente : {tx_tva:.1f}% vs taux attendu secteur {bench['tx_tva']:.0f}% "
                           f"— ecart {ecart_tva:.1f} pts. Verifier la coherence entre les taux appliques "
                           f"et le type d'activite (taux reduit, exoneration partielle...)."),
                'montant': 0, 'reference': '4457x'})

    return anomalies, resultats

def analyser_charges_salariales(df):
    """
    Analyse la coherence des charges de personnel :
    - Nombre de salaries estime vs charges
    - Rapport entre 421 (salaires) et 431 (cotisations)
    - Primes de fin d'annee suspectes
    """
    anomalies = []
    def s(pfx, col):
        return df[df['CompteNum'].str.startswith(pfx)][col].sum()

    sal_bruts   = s('641', 'Debit')
    cotis_pat   = s('645', 'Debit') + s('646', 'Debit')
    cotis_tot   = s('64', 'Debit') - sal_bruts
    avantages   = s('647', 'Debit')
    ca          = s('70', 'Credit') - s('70', 'Debit')

    # Ratio cotisations patronales / salaires bruts (attendu ~42-48%)
    ratio_cotis = cotis_pat / sal_bruts * 100 if sal_bruts > 0 else 0

    # Estimation nombre de salaries (SMIC ~1800 EUR/mois annuel ~21600 EUR)
    nb_sal_estime = round(sal_bruts / 25000) if sal_bruts > 0 else 0

    if sal_bruts > 0 and ratio_cotis > 0:
        if ratio_cotis < 20:
            anomalies.append({'type': 'Cotisations patronales trop faibles', 'gravite': 'ALERTE',
                'detail': (f"Cotisations/Salaires bruts : {ratio_cotis:.1f}% (attendu 42-48%). "
                           f"Possible sous-declaration des cotisations sociales ou salaries mal imputes."),
                'montant': float(cotis_pat), 'reference': '645x'})
        elif ratio_cotis > 60:
            anomalies.append({'type': 'Cotisations patronales anormalement elevees', 'gravite': 'ATTENTION',
                'detail': (f"Cotisations/Salaires bruts : {ratio_cotis:.1f}% (attendu 42-48%). "
                           f"Verifier la classification des postes en 64x."),
                'montant': float(cotis_pat), 'reference': '645x'})

    # Compte 421 (personnel — remuneration due) vs 641 (remuneration personnel)
    cpt_421_c = s('421', 'Credit')
    cpt_421_d = s('421', 'Debit')
    if sal_bruts > 0:
        ecart_421 = abs(sal_bruts - cpt_421_c)
        if ecart_421 > sal_bruts * 0.2 and cpt_421_c > 0:
            anomalies.append({'type': 'Incoherence salaires vs compte 421', 'gravite': 'ATTENTION',
                'detail': (f"641x (salaires bruts) : {sal_bruts:,.0f} EUR / "
                           f"421x (credit) : {cpt_421_c:,.0f} EUR — ecart {ecart_421:,.0f} EUR. "
                           f"Verifier les journaux de paie."),
                'montant': float(ecart_421), 'reference': '421/641'})

    return anomalies, {
        'sal_bruts':     float(sal_bruts),
        'cotis_pat':     float(cotis_pat),
        'ratio_cotis':   round(ratio_cotis, 1),
        'nb_sal_estime': nb_sal_estime,
        'avantages':     float(avantages),
    }

def analyser_scoring_fiscal(df):
    """
    Scoring fiscal : coherence IS, TVA, contributions.
    Signale les situations a risque fiscal.
    """
    anomalies = []
    def s(pfx, col):
        return df[df['CompteNum'].str.startswith(pfx)][col].sum()

    ca    = s('70', 'Credit') - s('70', 'Debit')
    tp    = s('7', 'Credit')  - s('7', 'Debit')
    tc    = s('6', 'Debit')   - s('6', 'Credit')
    res   = tp - tc
    is_d  = s('695', 'Debit') + s('6951', 'Debit')  # IS
    cfe_d = s('63511', 'Debit') + s('63512', 'Debit')  # CFE/CVAE
    tva_c = s('44571', 'Credit')
    tva_d = df[df['CompteNum'].str.startswith(('44566', '44562', '44563'))]['Debit'].sum()

    # Taux IS apparent
    tx_is = is_d / res * 100 if res > 0 else 0

    if res > 100000 and is_d == 0:
        anomalies.append({'type': 'IS absent malgre resultat positif', 'gravite': 'ALERTE',
            'detail': (f"Resultat positif de {res:,.0f} EUR mais aucun compte 695x (IS) detecte. "
                       f"Verifier si l'entreprise est soumise a l'IS et si la charge a bien ete provisionnee."),
            'montant': 0, 'reference': '695x'})
    elif res > 0 and is_d > 0:
        if tx_is > 35:
            anomalies.append({'type': 'Taux IS apparent eleve', 'gravite': 'ATTENTION',
                'detail': (f"IS / Resultat = {tx_is:.1f}% (seuil normal IS PME : 15-25%). "
                           f"Verifier les bases imposables et les eventuels rappels."),
                'montant': float(is_d), 'reference': '695x'})
        elif tx_is < 5 and res > 50000:
            anomalies.append({'type': 'Taux IS anormalement faible', 'gravite': 'ATTENTION',
                'detail': (f"IS / Resultat = {tx_is:.1f}% malgre un resultat de {res:,.0f} EUR. "
                           f"Verifier les credits d'impot (CIR, CICE residuel) ou les reports deficitaires."),
                'montant': float(is_d), 'reference': '695x'})

    # TVA : ratio TVA collectee / TVA deductible
    if tva_c > 0 and tva_d > tva_c * 2:
        anomalies.append({'type': 'Credit TVA structure suspect', 'gravite': 'ALERTE',
            'detail': (f"TVA deductible ({tva_d:,.0f} EUR) > 2x TVA collectee ({tva_c:,.0f} EUR). "
                       f"Un credit de TVA recurrent peut attirer un controle fiscal. "
                       f"Verifier la cohérence entre achats et ventes."),
            'montant': float(tva_d - tva_c), 'reference': '4456x/4457x'})

    return anomalies, {
        'is_d': float(is_d), 'tx_is': round(tx_is, 1),
        'tva_c': float(tva_c), 'tva_d': float(tva_d),
        'res': float(res),
        'cfe_d': float(cfe_d),
    }

# ══════════════════════════════════════════════════════════════
# GENERATEUR DE COMMENTAIRES (v3 + enrichissements v4)
# ══════════════════════════════════════════════════════════════

def generer_commentaires(df, anomalies, stats, secteur_code='INCONNU'):
    mg    = stats.get('marges', {})
    ben   = stats.get('benford', {})
    fe    = stats.get('fin_exercice', {})
    wd    = stats.get('weekend', {})
    dup   = stats.get('doublons', {})
    mr    = stats.get('montants_ronds', {})
    bench = stats.get('benchmark', {})
    sal   = stats.get('charges_salariales', {})
    fisc  = stats.get('scoring_fiscal', {})

    nc  = sum(1 for a in anomalies if a['gravite'] == 'CRITIQUE')
    na  = sum(1 for a in anomalies if a['gravite'] == 'ALERTE')
    nat = sum(1 for a in anomalies if a['gravite'] == 'ATTENTION')
    score = min(100, nc * 25 + na * 10 + nat * 3)

    commentaires = []

    # ── GLOBAL ──
    bench_label = bench.get('secteur_label', 'secteur non determine')
    if score == 0:
        intro = (f"L'analyse automatique n'a revele aucune anomalie significative. "
                 f"Secteur detecte : {bench_label}. Les equilibres sont respectes et "
                 f"aucun indicateur de fraude n'est declenche.")
        niveau = "Favorable"; couleur_global = "green"
    elif score < 25:
        intro = (f"{len(anomalies)} point(s) mineurs a verifier. Secteur : {bench_label}. "
                 f"Le dossier parait globalement sain.")
        niveau = "Faible"; couleur_global = "green"
    elif score < 50:
        intro = (f"{na} alerte(s) et {nat} point(s) d'attention. "
                 f"Secteur : {bench_label}. Certaines zones meritent un controle approfondi.")
        niveau = "Modere"; couleur_global = "orange"
    else:
        intro = (f"{nc} anomalie(s) critique(s) et {na} alerte(s). "
                 f"Secteur : {bench_label}. Ces anomalies requierent une attention immediate.")
        niveau = "Eleve"; couleur_global = "red"

    commentaires.append({'type': 'global', 'titre': 'Synthese generale',
                         'texte': intro, 'niveau': niveau, 'couleur': couleur_global})

    # ── EQUILIBRE ──
    eq_s = stats.get('equilibre', {})
    nb_deseq = eq_s.get('nb', 0)
    if nb_deseq > 0:
        commentaires.append({'type': 'specifique', 'page': 'equilibre', 'titre': 'Equilibre D/C',
            'texte': (f"{nb_deseq} ecriture(s) desequilibrees pour un ecart total de "
                      f"{eq_s.get('ecart_total', 0):,.2f} EUR. "
                      f"Peut signaler une importation incomplete ou une modification manuelle du FEC."),
            'couleur': 'red'})
    else:
        commentaires.append({'type': 'specifique', 'page': 'equilibre', 'titre': 'Equilibre D/C',
            'texte': "Toutes les ecritures sont parfaitement equilibrees.",
            'couleur': 'green'})

    # ── BENFORD ──
    bchi   = ben.get('chi2', 0)
    btotal = ben.get('total', 0)
    if btotal < 100:
        commentaires.append({'type': 'specifique', 'page': 'benford', 'titre': 'Loi de Benford',
            'texte': f"Non applicable : seulement {btotal} montants (minimum 100 requis).",
            'couleur': 'gray'})
    elif bchi > 15.5:
        commentaires.append({'type': 'specifique', 'page': 'benford', 'titre': 'Loi de Benford — ALERTE',
            'texte': (f"Chi2={bchi:.1f} (seuil 15.5) — distribution anormale. "
                      f"Causes possibles : montants inventes, saisies manuelles frequentes, forfaits. "
                      f"Revue des pieces justificatives recommandee pour les chiffres surrepresentes."),
            'couleur': 'orange'})
    else:
        commentaires.append({'type': 'specifique', 'page': 'benford', 'titre': 'Loi de Benford',
            'texte': (f"Distribution conforme (Chi2={bchi:.1f} < 15.5). "
                      f"Analyse sur {btotal} montants. Aucun signal de manipulation statistique."),
            'couleur': 'green'})

    # ── MARGES ──
    ca = mg.get('ca', 0); tm = mg.get('tm', 0); res = mg.get('res', 0); cp = mg.get('cp', 0)
    if ca > 0:
        txt = f"CA : {ca:,.0f} EUR — marge brute : {tm:.1f}%"
        if bench.get('tm_mediane'):
            ecart = tm - bench['tm_mediane']
            txt += (f" vs mediane sectorielle {bench['tm_mediane']:.0f}% "
                    f"({'+'if ecart>=0 else ''}{ecart:.1f} pts)")
        if tm < 0:
            txt += " — MARGE NEGATIVE : verifier completude des ventes et imputation des achats."
            c = 'red'
        elif bench.get('tm_low') and tm < bench.get('tm_low', 0):
            txt += f" — en dessous du bas de fourchette sectorielle ({bench.get('tm_low',0):.0f}%)."
            c = 'orange'
        elif bench.get('tm_high') and tm > bench.get('tm_high', 0):
            txt += f" — au-dessus du haut de fourchette. Verifier si des achats manquent."
            c = 'orange'
        else:
            txt += f". Resultat d'exploitation : {res:,.0f} EUR ({mg.get('tr',0):.1f}% du CA)."
            c = 'green'
        if ca > 0 and cp / ca > 0.5:
            txt += (f" Charges personnel ({cp:,.0f} EUR) = {cp/ca*100:.1f}% du CA — "
                    f"a comparer avec la mediane sectorielle ({bench.get('cp_ca_mediane', 30):.0f}%).")
        commentaires.append({'type': 'specifique', 'page': 'marges',
                              'titre': 'Marges et resultats', 'texte': txt, 'couleur': c})
    else:
        commentaires.append({'type': 'specifique', 'page': 'marges',
                              'titre': 'Marges et resultats',
                              'texte': "Aucun compte de ventes (70x) detecte ou solde nul.",
                              'couleur': 'gray'})

    # ── TVA ──
    tx_tva = mg.get('tx_tva', 0)
    if mg.get('tva_c', 0) > 0:
        if tx_tva > 25:
            txt_tva = (f"Taux TVA apparent ({tx_tva:.1f}%) anormalement eleve (max legal 20%). "
                       f"Verifier que la TVA collectee n'est pas doublee.")
            c_tva = 'red'
        elif tx_tva < 5 and tx_tva > 0:
            txt_tva = (f"Taux TVA apparent tres faible ({tx_tva:.1f}%). "
                       f"Verifier exoneration TVA ou saisies HT/TTC.")
            c_tva = 'orange'
        else:
            txt_tva = (f"Taux TVA apparent ({tx_tva:.1f}%) coherent. "
                       f"TVA collectee : {mg.get('tva_c',0):,.0f} EUR / "
                       f"TVA deductible : {mg.get('tva_d',0):,.0f} EUR.")
            c_tva = 'green'
        commentaires.append({'type': 'specifique', 'page': 'marges',
                              'titre': 'Coherence TVA', 'texte': txt_tva, 'couleur': c_tva})

    # ── BENCHMARK SECTORIEL ──
    if bench and bench.get('ca', 0) > 0:
        tm_e   = bench.get('tm_entreprise')
        tm_med = bench.get('tm_mediane', bench.get('tm_med', 0))
        note   = bench.get('note_sectorielle', '')
        secteur_lbl = bench.get('secteur_label', '')
        txt_b = f"Secteur identifie : {secteur_lbl}. {note}"
        if tm_e is not None:
            fourchette = f"[{bench.get('tm_low',0):.0f}%-{bench.get('tm_high',0):.0f}%]"
            ecart = tm_e - tm_med
            position = "dans la fourchette normale" if bench.get('tm_low',0) <= tm_e <= bench.get('tm_high',0) else ("en dessous" if tm_e < bench.get('tm_low',0) else "au-dessus")
            txt_b += (f" Marge brute de l'entreprise : {tm_e:.1f}% — mediane secteur : {tm_med:.0f}% {fourchette} — "
                      f"position : {position} ({'+'if ecart>=0 else ''}{ecart:.1f} pts).")
        commentaires.append({'type': 'specifique', 'page': 'benchmark',
                              'titre': 'Analyse sectorielle', 'texte': txt_b,
                              'couleur': 'green' if position == "dans la fourchette normale" else 'orange'})

    # ── CHARGES SALARIALES ──
    if sal.get('sal_bruts', 0) > 0:
        rc = sal.get('ratio_cotis', 0)
        nb_sal = sal.get('nb_sal_estime', 0)
        if rc < 20 or rc > 60:
            c_sal = 'orange'
            txt_sal = (f"Ratio cotisations/salaires bruts : {rc:.1f}% "
                       f"(attendu 42-48%). Anomalie detectee — verifier les journaux de paie.")
        else:
            c_sal = 'green'
            txt_sal = (f"Ratio cotisations/salaires bruts : {rc:.1f}% (norme 42-48%) — coherent. "
                       f"Effectif estime : {nb_sal} salarie(s) "
                       f"(base salaire moyen 25 000 EUR/an).")
        commentaires.append({'type': 'specifique', 'page': 'benchmark',
                              'titre': 'Charges salariales', 'texte': txt_sal, 'couleur': c_sal})

    # ── SCORING FISCAL ──
    if fisc.get('res', 0) > 0:
        tx_is = fisc.get('tx_is', 0)
        is_d  = fisc.get('is_d', 0)
        if is_d == 0 and fisc.get('res', 0) > 100000:
            txt_f = "Resultat positif mais aucun IS detecte — verifier la provision pour IS."
            c_f   = 'orange'
        elif tx_is > 35:
            txt_f = f"Taux IS apparent eleve ({tx_is:.1f}%) — verifier les bases imposables."
            c_f   = 'orange'
        else:
            txt_f = (f"Taux IS apparent : {tx_is:.1f}% — dans les normes. "
                     f"IS comptabilise : {is_d:,.0f} EUR.")
            c_f   = 'green'
        commentaires.append({'type': 'specifique', 'page': 'benchmark',
                              'titre': 'Position fiscale', 'texte': txt_f, 'couleur': c_f})

    # ── DOUBLONS ──
    nb_dup = dup.get('nb', 0)
    if nb_dup > 0:
        commentaires.append({'type': 'specifique', 'page': 'doublons', 'titre': 'Doublons detectes',
            'texte': (f"{nb_dup} lignes en doublon exact (meme date, compte, montant). "
                      f"Double importation, correction mal effectuee ou fraude possible. "
                      f"Chaque doublon doit etre justifie par une piece distincte."),
            'couleur': 'orange'})
    else:
        commentaires.append({'type': 'specifique', 'page': 'doublons', 'titre': 'Doublons',
            'texte': "Aucun doublon exact detecte.", 'couleur': 'green'})

    # ── WEEK-END ──
    nb_we = wd.get('nb', 0); pct_we = wd.get('pct', 0)
    if nb_we > 0:
        if pct_we > 10:
            txt_we = (f"{nb_we} ecritures week-end ({pct_we:.1f}%). Selon l'activite "
                      f"({bench.get('secteur_label','')}), cela peut etre normal (commerce, restauration). "
                      f"Pour du B2B ou conseil, justifier systematiquement.")
            c_we = 'orange'
        else:
            txt_we = (f"{nb_we} ecritures week-end ({pct_we:.1f}%) — taux faible, "
                      f"probablement des regularisations ou imports automatiques.")
            c_we = 'gray'
    else:
        txt_we = "Aucune ecriture le week-end. Saisie concentree sur les jours ouvrables."
        c_we = 'green'
    commentaires.append({'type': 'specifique', 'page': 'weekend',
                         'titre': 'Ecritures hors jours ouvres', 'texte': txt_we, 'couleur': c_we})

    # ── FIN D'EXERCICE ──
    p7 = fe.get('p7', 0)
    if p7 > 25:
        commentaires.append({'type': 'specifique', 'page': 'cloture',
                              'titre': 'Concentration fin exercice',
            'texte': (f"{p7:.1f}% du volume sur les 7 derniers jours. "
                      f"'Habillage de bilan' possible : provisions massives, anticipation de produits, "
                      f"ecritures fictives. Examiner chaque ecriture de cloture."),
            'couleur': 'red'})
    elif p7 > 10:
        commentaires.append({'type': 'specifique', 'page': 'cloture', 'titre': 'Fin exercice',
            'texte': f"{p7:.1f}% du volume sur les 7 derniers jours — taux modere, a surveiller.",
            'couleur': 'orange'})
    else:
        commentaires.append({'type': 'specifique', 'page': 'cloture', 'titre': 'Fin exercice',
            'texte': f"Volume de fin d'exercice normal ({p7:.1f}% sur les 7 derniers jours).",
            'couleur': 'green'})

    # ── MONTANTS RONDS ──
    p1000 = mr.get('p1000', 0)
    if p1000 > 20:
        commentaires.append({'type': 'specifique', 'page': 'montants', 'titre': 'Montants ronds',
            'texte': (f"{p1000:.1f}% des ecritures sont des multiples de 1000 EUR. "
                      f"Estimations comptables, charges fictives ou saisies forfaitaires possibles. "
                      f"Demander les justificatifs pour les plus significatifs."),
            'couleur': 'orange'})

    # ── BOUCLIER FISCAL (v5) ──
    bouclier_s = stats.get('bouclier_fiscal', {})
    b_score_c  = bouclier_s.get('score_global', 0)
    b_niveau_c = bouclier_s.get('niveau_risque', '')
    b_conseil_c = bouclier_s.get('conseil', '')
    if bouclier_s:
        if b_score_c >= 60:
            c_b = 'red'
            txt_b2 = (f"Score de probabilite de controle DGFiP : {b_score_c}/100 — {b_niveau_c}. "
                      f"{b_conseil_c} "
                      f"Les axes les plus a risque : "
                      + ', '.join(
                          v['label'] for v in bouclier_s.get('details', {}).values()
                          if v.get('points', 0) > 0
                      ) + '.')
        elif b_score_c >= 40:
            c_b = 'orange'
            txt_b2 = (f"Score Bouclier Fiscal : {b_score_c}/100 — {b_niveau_c}. "
                      f"{b_conseil_c}")
        elif b_score_c >= 20:
            c_b = 'orange'
            txt_b2 = (f"Score Bouclier Fiscal : {b_score_c}/100 — {b_niveau_c}. "
                      f"Quelques signaux mineurs a documenter.")
        else:
            c_b = 'green'
            txt_b2 = (f"Score Bouclier Fiscal : {b_score_c}/100 — {b_niveau_c}. "
                      f"Aucun signal de risque fiscal majeur detecte.")
        commentaires.append({'type': 'specifique', 'page': 'bouclier',
                             'titre': f'Bouclier Fiscal — Score {b_score_c}/100',
                             'texte': txt_b2, 'couleur': c_b})

    # ── SUGGESTIONS ──
    suggestions = []
    if mg.get('ca', 0) > 0 and mg.get('tva_c', 0) == 0:
        suggestions.append("Aucun compte TVA collectee (44571x) detecte malgre un CA — verifier le regime TVA.")
    if stats.get('equilibre', {}).get('nb', 0) > 5:
        suggestions.append("Nombre eleve de desequilibres D/C — reimporter depuis la source plutot que corriger manuellement.")
    if mg.get('cp', 0) > 0 and mg.get('ca', 0) == 0:
        suggestions.append("Charges de personnel sans ventes — exercice partiel ou comptes CA manquants ?")
    if ben.get('total', 0) > 1000 and ben.get('chi2', 0) > 20:
        suggestions.append("Score Benford tres eleve — auditer les pieces pour les chiffres surrepresentes.")
    if dup.get('nb', 0) > 10:
        suggestions.append("Nombre important de doublons — verifier si le FEC a ete genere deux fois.")
    if bench.get('tm_entreprise') is not None and bench.get('tm_low'):
        if bench['tm_entreprise'] < bench['tm_low']:
            suggestions.append(
                f"Marge brute ({bench['tm_entreprise']:.1f}%) sous la fourchette sectorielle "
                f"({bench['tm_low']:.0f}%-{bench['tm_high']:.0f}%) — risque de dissimulation de recettes.")
    if sal.get('nb_sal_estime', 0) > 0:
        suggestions.append(
            f"Effectif estime : {sal['nb_sal_estime']} salarie(s). "
            f"Croiser avec les declarations DPAE/DSN si disponibles.")
    if b_score_c >= 60:
        suggestions.append(
            f"Bouclier Fiscal ELEVE ({b_score_c}/100) — recommander une mission de pre-audit fiscal "
            f"avant la prochaine echeance de controle.")
    elif b_score_c >= 40:
        suggestions.append(
            f"Bouclier Fiscal MODERE ({b_score_c}/100) — renforcer la documentation "
            f"des zones identifiees avant depot des declarations.")

    return commentaires, suggestions

# ══════════════════════════════════════════════════════════════
# RAPPORT HTML v4
# ══════════════════════════════════════════════════════════════

def g(val, ok_min=None, ok_max=None):
    if ok_min is not None and val < ok_min: return 'var(--red)'
    if ok_max is not None and val > ok_max: return 'var(--red)'
    return 'var(--green)'

def generer_rapport(filepath, df, anomalies, stats, commentaires, suggestions, secteur_code='INCONNU'):
    mg    = stats.get('marges', {})
    ben   = stats.get('benford', {})
    fe    = stats.get('fin_exercice', {})
    wd    = stats.get('weekend', {})
    mr    = stats.get('montants_ronds', {})
    sais  = stats.get('saisonnalite', {})
    conc  = stats.get('concentration', {})
    bench = stats.get('benchmark', {})
    sal   = stats.get('charges_salariales', {})
    fisc  = stats.get('scoring_fiscal', {})

    c     = Counter(a['gravite'] for a in anomalies)
    nc, na, nat, ni = c.get('CRITIQUE',0), c.get('ALERTE',0), c.get('ATTENTION',0), c.get('INFO',0)
    score = min(100, nc*25 + na*10 + nat*3 + ni)
    sc    = '#ef4444' if score >= 50 else '#f97316' if score >= 25 else '#22c55e'
    sl    = 'ELEVE' if score >= 50 else 'MODERE' if score >= 25 else 'FAIBLE'

    # Graphiques (données)
    vol      = sais.get('vol_mensuel', [])
    has_mois = len(vol) >= 2
    lm = json.dumps([str(v[0]) for v in vol]) if has_mois else '[]'
    dm = json.dumps([v[1] for v in vol]) if has_mois else '[]'

    cam     = sais.get('ca_mensuel', {})
    has_ca  = len(cam) >= 2 and any(v > 0 for v in cam.values())
    lca     = json.dumps(list(cam.keys())) if has_ca else '[]'
    dca     = json.dumps([round(v/1000, 1) for v in cam.values()]) if has_ca else '[]'

    pj        = wd.get('par_jour', {})
    has_jours = len(pj) >= 2
    ljours    = json.dumps(list(pj.keys())) if has_jours else '[]'
    djours    = json.dumps(list(pj.values())) if has_jours else '[]'

    has_gravite = (nc + na + nat + ni) > 0

    sol_data = stats.get('soldes', {}).get('soldes', [])
    cls_d = defaultdict(float); cls_c = defaultdict(float)
    for s2 in sol_data:
        cl = s2.get('CompteNum', '?')[:1]
        cls_d[cl] += float(s2.get('debit', 0))
        cls_c[cl] += float(s2.get('credit', 0))
    all_cl    = sorted(set(list(cls_d.keys()) + list(cls_c.keys())))
    has_classes = len(all_cl) >= 2
    cl_labels   = json.dumps(all_cl) if has_classes else '[]'
    cl_d        = json.dumps([round(cls_d.get(k, 0)/1000, 1) for k in all_cl]) if has_classes else '[]'
    cl_c        = json.dumps([round(cls_c.get(k, 0)/1000, 1) for k in all_cl]) if has_classes else '[]'

    ben_obs  = ben.get('obs', {}); ben_att = ben.get('att', {})
    has_benford = len(ben_obs) == 9 and ben.get('total', 0) >= 100
    b_labels = json.dumps([str(i) for i in range(1, 10)])
    b_obs    = json.dumps([ben_obs.get(str(i), 0) for i in range(1, 10)])
    b_att    = json.dumps([ben_att.get(str(i), 0) for i in range(1, 10)])

    tc       = conc.get('top_clients', [])
    has_conc = len(tc) >= 2
    tc_json  = json.dumps([{'lib': x.get('CompAuxLib', '?'), 'pct': float(x.get('pct', 0))} for x in tc[:5]])

    p7     = fe.get('p7', 0); p_rest = max(0, 100 - p7)
    has_fe = (p7 + p_rest) > 0

    # Données benchmark pour graphe radar
    tm_e   = bench.get('tm_entreprise', 0) or 0
    tm_med = bench.get('tm_mediane', bench.get('tm_med', 0)) or 0
    tm_low = bench.get('tm_low', 0) or 0
    tm_high= bench.get('tm_high', 0) or 0
    cp_e   = bench.get('cp_ca_entreprise', 0) or 0
    cp_med = bench.get('cp_ca_mediane', 0) or 0
    has_bench = bench.get('ca', 0) > 0

    # ── Fonctions helpers ──
    def rows(lst):
        if not lst:
            return '<tr><td colspan="5" class="empty-row">Aucune anomalie pour ce controle</td></tr>'
        r = ""
        for a in lst:
            bg = {'CRITIQUE':'#fef2f2','ALERTE':'#fff7ed','ATTENTION':'#fffbeb','INFO':'#f0fdf4'}.get(a['gravite'],'#f9fafb')
            bc = {'CRITIQUE':'#dc2626','ALERTE':'#ea580c','ATTENTION':'#ca8a04','INFO':'#16a34a'}.get(a['gravite'],'#6b7280')
            r += (f'<tr style="background:{bg}"><td><span class="badge" style="background:{bc}">'
                  f'{a["gravite"]}</span></td><td class="t-type">{a["type"]}</td>'
                  f'<td class="t-det">{a["detail"]}</td><td class="t-amt">{a["montant"]:,.0f} EUR</td>'
                  f'<td class="t-ref">{a["reference"]}</td></tr>')
        return r

    def nbadge(lst, col='red'):
        if not lst: return ''
        colors = {'red':'#ef4444','orange':'#f97316','yellow':'#ca8a04','green':'#22c55e'}
        return f'<span class="nbadge" style="background:{colors.get(col,"#ef4444")}">{len(lst)}</span>'

    def comment_html(page=None):
        coms = [cm for cm in commentaires
                if (page is None and cm['type'] == 'global') or (page and cm.get('page') == page)]
        if not coms: return ''
        html2 = ''
        for cm in coms:
            cc = {'red':'rgba(239,68,68,.08)','orange':'rgba(249,115,22,.08)',
                  'green':'rgba(34,197,94,.08)','gray':'rgba(100,100,130,.08)'}.get(cm['couleur'],'rgba(100,100,130,.08)')
            bc = {'red':'#ef4444','orange':'#f97316','green':'#22c55e','gray':'#6b7280'}.get(cm['couleur'],'#6b7280')
            icon = {'red':'⚠️','orange':'⚡','green':'✅','gray':'ℹ️'}.get(cm['couleur'],'ℹ️')
            html2 += (f'<div class="com-box" style="background:{cc};border-left-color:{bc}">'
                      f'<div class="com-title">{icon} {cm["titre"]}</div>'
                      f'<div class="com-text">{cm["texte"]}</div></div>')
        return html2

    global_com = next((cm for cm in commentaires if cm['type'] == 'global'), None)
    gcom_html = ''
    if global_com:
        cc = {'red':'rgba(239,68,68,.08)','orange':'rgba(249,115,22,.08)',
              'green':'rgba(34,197,94,.08)'}.get(global_com['couleur'],'rgba(100,100,130,.08)')
        bc = {'red':'#ef4444','orange':'#f97316','green':'#22c55e'}.get(global_com['couleur'],'#6b7280')
        gcom_html = (f'<div class="com-box" style="background:{cc};border-left-color:{bc};margin-bottom:20px">'
                     f'<div class="com-title">📋 {global_com["titre"]} — Risque {global_com["niveau"]}</div>'
                     f'<div class="com-text">{global_com["texte"]}</div></div>')

    sugg_html = ''
    if suggestions:
        items = ''.join(f'<li>{s2}</li>' for s2 in suggestions)
        sugg_html = (f'<div class="com-box" style="background:rgba(155,138,251,.08);'
                     f'border-left-color:#9b8afb;margin-top:16px">'
                     f'<div class="com-title">💡 Suggestions de l\'analyse automatique</div>'
                     f'<ul class="sugg-list">{items}</ul></div>')

    def chart_box(chart_id, title, has_data, no_data_msg, height=180, extra=''):
        if not has_data:
            return f'<div class="cb"><h4>{title}</h4><div class="no-chart">{no_data_msg}</div></div>'
        return f'<div class="cb"><h4>{title}</h4>{extra}<canvas id="{chart_id}" height="{height}"></canvas></div>'

    # Catégories anomalies
    cats = {
        'equilibre':    [a for a in anomalies if 'Desequilibre' in a['type']],
        'benford':      [a for a in anomalies if 'Benford' in a['type']],
        'ronds':        [a for a in anomalies if 'rond' in a['type'].lower() or 'aberrant' in a['type'].lower()],
        'weekend':      [a for a in anomalies if 'week-end' in a['type'].lower() or 'Week' in a['type']],
        'soldes':       [a for a in anomalies if any(x in a['type'] for x in ['Charge c','Produit d','Fournisseur d','Client c'])],
        'doublons':     [a for a in anomalies if any(x in a['type'] for x in ['Doublon','Inverse','Extourne'])],
        'cloture':      [a for a in anomalies if any(x in a['type'] for x in ['Concentration','Trous'])],
        'concentration':[a for a in anomalies if 'dominant' in a['type'].lower()],
        'marges':       [a for a in anomalies if any(x in a['type'] for x in ['Marge','Personnel','TVA inco','TVA ded'])],
        'saisonnalite': [a for a in anomalies if 'regulier' in a['type'].lower()],
        'benchmark':    [a for a in anomalies if 'mediane' in a['type'].lower() or
                         'sectoriel' in a['type'].lower() or
                         'cotisations' in a['type'].lower() or 'IS' in a['type'] or
                         'inattendue' in a['type'].lower()],
    }

    # Marge items HTML
    def mi(label, val, color='var(--text)'):
        return f'<div class="mi"><div class="ml">{label}</div><div class="mv" style="color:{color}">{val}</div></div>'

    marge_items = (
        mi("Chiffre d'affaires (70x)", f"{mg.get('ca',0):,.2f} EUR", "var(--cyan)") +
        mi("Achats (60x)", f"{mg.get('achats',0):,.2f} EUR", "var(--orange)") +
        mi("Marge brute", f"{mg.get('mb',0):,.2f} EUR", "var(--green)" if mg.get('mb',0)>=0 else "var(--red)") +
        mi("Taux marge brute", f"{mg.get('tm',0):.2f} %", "var(--green)" if mg.get('tm',0)>0 else "var(--red)") +
        mi("Mediane sectorielle", f"{tm_med:.1f} % [{tm_low:.0f}–{tm_high:.0f}%]",
           "var(--violet)") +
        mi("Charges externes (61+62)", f"{mg.get('ce',0):,.2f} EUR") +
        mi("Charges de personnel (64x)", f"{mg.get('cp',0):,.2f} EUR", "var(--yellow)") +
        mi("Total produits (7xx)", f"{mg.get('tp',0):,.2f} EUR") +
        mi("Total charges (6xx)", f"{mg.get('tc',0):,.2f} EUR") +
        mi("Resultat d'exploitation", f"{mg.get('res',0):,.2f} EUR", "var(--green)" if mg.get('res',0)>=0 else "var(--red)") +
        mi("Taux de resultat", f"{mg.get('tr',0):.2f} %") +
        mi("TVA collectee (4457x)", f"{mg.get('tva_c',0):,.2f} EUR") +
        mi("TVA deductible (4456x)", f"{mg.get('tva_d',0):,.2f} EUR") +
        mi("Taux TVA apparent", f"{mg.get('tx_tva',0):.2f} %",
           "var(--red)" if mg.get('tx_tva',0) > 20 else "var(--text)")
    )

    # Benchmark items HTML
    bench_items = ""
    if has_bench:
        ecart_tm = round(tm_e - tm_med, 1) if tm_e and tm_med else 0
        ecart_color_tm = "var(--red)" if (tm_e < tm_low or tm_e > tm_high) else "var(--green)"
        ecart_cp = round(cp_e - cp_med, 1) if cp_e and cp_med else 0
        bench_items = (
            mi("Secteur identifie", bench.get('secteur_label','N/A'), "var(--violet)") +
            mi("Marge brute entreprise", f"{tm_e:.1f}%", "var(--cyan)") +
            mi("Mediane sectorielle", f"{tm_med:.0f}% [{tm_low:.0f}–{tm_high:.0f}%]", "var(--muted)") +
            mi("Ecart vs mediane", f"{'+'if ecart_tm>=0 else ''}{ecart_tm:.1f} pts", ecart_color_tm) +
            mi("Charges pers./CA entreprise", f"{cp_e:.1f}%", "var(--yellow)") +
            mi("Charges pers./CA secteur", f"{cp_med:.0f}%", "var(--muted)") +
            mi("Ecart charges pers.", f"{'+'if ecart_cp>=0 else ''}{ecart_cp:.1f} pts",
               "var(--red)" if abs(ecart_cp) > 15 else "var(--green)") +
            mi("TVA attendue secteur", f"{bench.get('tx_tva_attendu',0):.0f}%", "var(--muted)") +
            mi("TVA observee", f"{bench.get('tx_tva_observe',0):.1f}%", "var(--text)") +
            mi("Salaires bruts", f"{sal.get('sal_bruts',0):,.0f} EUR", "var(--text)") +
            mi("Cotis. patronales", f"{sal.get('cotis_pat',0):,.0f} EUR", "var(--text)") +
            mi("Ratio cotis./salaires", f"{sal.get('ratio_cotis',0):.1f}%",
               "var(--red)" if not (20 <= sal.get('ratio_cotis',0) <= 60) else "var(--green)") +
            mi("Effectif estime", f"~{sal.get('nb_sal_estime',0)} salarie(s)", "var(--text)") +
            mi("IS comptabilise", f"{fisc.get('is_d',0):,.0f} EUR", "var(--text)") +
            mi("Taux IS apparent", f"{fisc.get('tx_is',0):.1f}%",
               "var(--red)" if fisc.get('tx_is',0) > 35 else "var(--text)")
        )

    # Top clients
    top_html = ""
    for item in tc[:5]:
        pct = float(item.get('pct', 0))
        top_html += (f'<div class="cr"><div class="cn">{item.get("CompAuxLib","?")} '
                     f'<span class="cnum">{item.get("CompAuxNum","")}</span></div>'
                     f'<div class="cbar"><div class="cbf" style="width:{min(100,pct)}%"></div></div>'
                     f'<div class="cp">{pct:.1f}%</div></div>')
    if not top_html:
        top_html = '<p class="nd">Donnees tiers non disponibles (CompAuxNum absent)</p>'

    # Données jauge benchmark pour Chart.js
    bench_chart_labels = json.dumps(['Marge entreprise', 'Mediane secteur', 'Bas fourchette', 'Haut fourchette'])
    bench_chart_data   = json.dumps([round(tm_e,1), round(tm_med,1), round(tm_low,1), round(tm_high,1)])
    cp_chart_labels    = json.dumps(["Pers./CA entreprise", "Pers./CA secteur"])
    cp_chart_data      = json.dumps([round(cp_e,1), round(cp_med,1)])

    # ── Données Bouclier Fiscal ──
    bouclier      = stats.get('bouclier_fiscal', {})
    b_score       = bouclier.get('score_global', 0)
    b_niveau      = bouclier.get('niveau_risque', 'N/A')
    b_conseil     = bouclier.get('conseil', '')
    b_details     = bouclier.get('details', {})
    b_max_th      = bouclier.get('max_theorique', 105)
    # Couleur score bouclier
    b_color = ('#ef4444' if b_score >= 80 else '#f97316' if b_score >= 60
               else '#eab308' if b_score >= 40 else '#22c55e')
    b_dash  = round(b_score / 100 * 339.3, 1)  # circumference r=54 ≈ 339.3

    def axe_html(key, info):
        pts   = info.get('points', 0)
        mx    = info.get('max', 1)
        lbl   = info.get('label', key)
        det   = info.get('detail', '')
        pct_w = round(pts / mx * 100)
        color = ('#ef4444' if pct_w >= 80 else '#f97316' if pct_w >= 50
                 else '#eab308' if pct_w >= 25 else '#06d6a0')
        pts_color = color if pts > 0 else '#6e6e90'
        return (f'<div class="axe-card">'
                f'<div class="axe-header">'
                f'<span class="axe-label">{lbl}</span>'
                f'<span class="axe-pts" style="color:{pts_color}">{pts}/{mx}</span>'
                f'</div>'
                f'<div class="axe-bar-track">'
                f'<div class="axe-bar-fill" style="width:{pct_w}%;background:{color}"></div>'
                f'</div>'
                f'<div class="axe-detail">{det}</div>'
                f'</div>')

    axes_html = ''.join(axe_html(k, v) for k, v in b_details.items())

    # ── Données Prévisionnel N+1 ──
    prev         = stats.get('previsionnel', {})
    has_prev     = bool(prev and prev.get('ca_n1', 0) > 0)
    p_annee_n    = prev.get('annee_n', 'N') if has_prev else 'N'
    try:
        p_annee_n1 = str(int(p_annee_n) + 1)
    except:
        p_annee_n1 = 'N+1'
    p_croiss     = prev.get('croissance_ca_hyp', 0.05) if has_prev else 0.05
    p_ca_n       = prev.get('ca_n', 0)
    p_ca_n1      = prev.get('ca_n1', 0)
    p_ach_n      = prev.get('achats_n', 0)
    p_ach_n1     = prev.get('achats_n1', 0)
    p_mb_n       = p_ca_n - p_ach_n
    p_mb_n1      = prev.get('mb_n1', 0)
    p_cp_n       = prev.get('cp_n', 0)
    p_cp_n1      = prev.get('cp_n1', 0)
    p_ce_n       = prev.get('ce_n', 0)
    p_ce_n1      = prev.get('ce_n1', 0)
    p_amo_n      = prev.get('dot_amo_n', 0)
    p_amo_n1     = prev.get('dot_amo_n1', 0)
    p_ebe_n1     = prev.get('ebe_n1', 0)
    p_rex_n      = prev.get('res_n', 0)
    p_rex_n1     = prev.get('rex_n1', 0)
    p_is_n1      = prev.get('is_n1', 0)
    p_res_n1     = prev.get('res_net_n1', 0)
    p_tm_n1      = prev.get('tm_n1', 0)
    p_tr_n1      = prev.get('tr_n1', 0)
    p_cp_ca_n1   = prev.get('cp_ca_n1', 0)
    p_nb_sal     = prev.get('nb_sal_estime', 0)
    p_taux_is    = prev.get('taux_is', 0.25)

    def pvar(n, n1):
        """Formatte la variation avec couleur."""
        if n == 0: return '<span style="color:var(--muted)">—</span>'
        v = n1 - n; p = v / abs(n) * 100
        cls = 'var-pos' if v >= 0 else 'var-neg'
        return f'<span class="{cls}">{"+"if v>=0 else ""}{v:,.0f} ({p:+.1f}%)</span>'

    def pfmt(v, color=''):
        s = f'{v:,.0f}'
        if color: return f'<span style="color:{color}">{s}</span>'
        return s

    # CA mensuel N+1 pour graphique
    cam_n1    = prev.get('ca_mensuel_n1', {}) if has_prev else {}
    has_prev_ca = len(cam_n1) >= 2
    lca_n1    = json.dumps(list(cam_n1.keys()))
    dca_n1    = json.dumps([round(v/1000, 1) for v in cam_n1.values()])
    cam_n_prev = prev.get('ca_mensuel_n', stats.get('saisonnalite', {}).get('ca_mensuel', {}))
    dca_n_prev = json.dumps([round(cam_n_prev.get(k, 0)/1000, 1) for k in cam_n1.keys()]) if has_prev_ca else '[]'

    # ── HTML Previsionnel pre-rendu (evite les f-strings imbriques) ──
    prev_no_data_html = (
        '<div class="com-box" style="background:rgba(255,209,102,.08);border-left-color:#eab308">'
        '<div class="com-title">&#x26A0;&#xFE0F; Donnees insuffisantes</div>'
        '<div class="com-text">Aucune donnee de ventes detectee dans le FEC. '
        'Le previsionnel ne peut pas etre genere.</div></div>'
    ) if not has_prev else ''

    if has_prev:
        _mb_color  = 'var(--green)' if p_mb_n1  >= 0 else 'var(--red)'
        _ebe_color = 'var(--green)' if p_ebe_n1 >= 0 else 'var(--red)'
        _res_color = 'var(--green)' if p_res_n1 >= 0 else 'var(--red)'
        _tr_color  = 'var(--green)' if p_tr_n1  >= 0 else 'var(--red)'
        _chart_ca  = chart_box('cPrevCA',
                               f'CA mensuel {p_annee_n} vs {p_annee_n1} (k EUR)',
                               has_prev_ca, 'Saisonnalite mensuelle non disponible', 200)
        prev_content_html = (
            f'<div class="hyp-box">'
            f'<h4>Hypotheses de projection</h4>'
            f'<div class="hyp-grid">'
            f'<div class="hyp-item">Exercice de reference : <b>{p_annee_n}</b></div>'
            f'<div class="hyp-item">Exercice projete : <b>{p_annee_n1}</b></div>'
            f'<div class="hyp-item">Croissance CA retenue : <b>+{p_croiss*100:.1f}%</b></div>'
            f'<div class="hyp-item">Revalorisation masse salariale : <b>+3,0%</b></div>'
            f'<div class="hyp-item">Amortissements : <b>identiques a N</b></div>'
            f'<div class="hyp-item">Charges externes : <b>proportionnelles au CA</b></div>'
            f'<div class="hyp-item">Taux IS retenu : <b>{p_taux_is*100:.0f}%</b></div>'
            f'<div class="hyp-item">Effectif estime : <b>~{p_nb_sal} salarie(s)</b></div>'
            f'</div></div>'

            f'<div class="prev-kgrid">'
            f'<div class="kpi"><div class="kl">CA previsionnel {p_annee_n1}</div>'
            f'<div class="kv" style="color:var(--cyan)">{p_ca_n1:,.0f}</div><div class="ks">EUR</div></div>'
            f'<div class="kpi"><div class="kl">Marge brute N+1</div>'
            f'<div class="kv" style="color:{_mb_color}">{p_mb_n1:,.0f}</div>'
            f'<div class="ks">Taux {p_tm_n1:.1f}%</div></div>'
            f'<div class="kpi"><div class="kl">EBE previsionnel</div>'
            f'<div class="kv" style="color:{_ebe_color}">{p_ebe_n1:,.0f}</div>'
            f'<div class="ks">EUR</div></div>'
            f'<div class="kpi"><div class="kl">Resultat net N+1</div>'
            f'<div class="kv" style="color:{_res_color}">{p_res_n1:,.0f}</div>'
            f'<div class="ks">Taux {p_tr_n1:.1f}%</div></div>'
            f'<div class="kpi"><div class="kl">Charges pers./CA</div>'
            f'<div class="kv">{p_cp_ca_n1:.1f}%</div>'
            f'<div class="ks">vs sect. {bench.get("cp_ca_mediane",0):.0f}%</div></div>'
            f'<div class="kpi"><div class="kl">IS previsionnel</div>'
            f'<div class="kv">{p_is_n1:,.0f}</div>'
            f'<div class="ks">Taux {p_taux_is*100:.0f}%</div></div>'
            f'</div>'

            f'<div class="sec"><div class="sh"><span>&#x1F4D0;</span>'
            f'<h3>Compte de resultat {p_annee_n} / {p_annee_n1}</h3>'
            f'<span class="sd">Toutes les valeurs en EUR</span></div>'
            f'<div class="sb0"><table class="prev-table"><thead><tr>'
            f'<th style="text-align:left">Poste</th>'
            f'<th>{p_annee_n} (reel)</th><th>{p_annee_n1} (prev.)</th><th>Variation</th>'
            f'</tr></thead><tbody>'
            f'<tr><td>Chiffre d\'affaires</td>'
            f'<td>{pfmt(p_ca_n,"var(--cyan)")}</td><td>{pfmt(p_ca_n1,"var(--cyan)")}</td>'
            f'<td>{pvar(p_ca_n,p_ca_n1)}</td></tr>'
            f'<tr class="ded"><td>&nbsp;&nbsp;&mdash; Achats / CAMV</td>'
            f'<td>{pfmt(p_ach_n)}</td><td>{pfmt(p_ach_n1)}</td><td>{pvar(p_ach_n,p_ach_n1)}</td></tr>'
            f'<tr class="subtotal"><td>= Marge brute</td>'
            f'<td>{pfmt(p_mb_n)}</td><td>{pfmt(p_mb_n1)}</td><td>{pvar(p_mb_n,p_mb_n1)}</td></tr>'
            f'<tr class="ded"><td>&nbsp;&nbsp;&mdash; Charges externes</td>'
            f'<td>{pfmt(p_ce_n)}</td><td>{pfmt(p_ce_n1)}</td><td>{pvar(p_ce_n,p_ce_n1)}</td></tr>'
            f'<tr class="ded"><td>&nbsp;&nbsp;&mdash; Charges de personnel</td>'
            f'<td>{pfmt(p_cp_n)}</td><td>{pfmt(p_cp_n1)}</td><td>{pvar(p_cp_n,p_cp_n1)}</td></tr>'
            f'<tr class="subtotal"><td>= EBE</td>'
            f'<td>&mdash;</td><td>{pfmt(p_ebe_n1,"var(--cyan)")}</td><td>&mdash;</td></tr>'
            f'<tr class="ded"><td>&nbsp;&nbsp;&mdash; Dotations amort.</td>'
            f'<td>{pfmt(p_amo_n)}</td><td>{pfmt(p_amo_n1)}</td><td>{pvar(p_amo_n,p_amo_n1)}</td></tr>'
            f'<tr class="subtotal"><td>= Resultat d\'exploitation</td>'
            f'<td>{pfmt(p_rex_n)}</td><td>{pfmt(p_rex_n1)}</td><td>{pvar(p_rex_n,p_rex_n1)}</td></tr>'
            f'<tr class="ded"><td>&nbsp;&nbsp;&mdash; IS ({p_taux_is*100:.0f}%)</td>'
            f'<td>&mdash;</td><td>{pfmt(p_is_n1)}</td><td>&mdash;</td></tr>'
            f'<tr class="total"><td>= Resultat net</td>'
            f'<td>{pfmt(p_rex_n)}</td>'
            f'<td>{pfmt(p_res_n1,"var(--green)" if p_res_n1>=0 else "var(--red)")}</td>'
            f'<td>{pvar(p_rex_n,p_res_n1)}</td></tr>'
            f'</tbody></table></div></div>'

            f'<div class="g2">'
            f'{_chart_ca}'
            f'<div class="cb"><h4>Ratios cles N+1</h4><div style="margin-top:8px">'
            f'{mi("Taux marge brute", f"{p_tm_n1:.1f}%", "var(--cyan)")}'
            f'{mi("Taux resultat net", f"{p_tr_n1:.1f}%", _tr_color)}'
            f'{mi("Charges personnel/CA", f"{p_cp_ca_n1:.1f}%", "var(--yellow)")}'
            + mi('Secteur median charges', str(round(bench.get('cp_ca_mediane',0))) + '%', 'var(--muted)') +
            f'{mi("IS taux retenu", f"{p_taux_is*100:.0f}%", "var(--muted)")}'
            f'{mi("Effectif estime", f"~{p_nb_sal} salarie(s)", "var(--text)")}'
            f'</div></div></div>'

            f'<div class="warn-prev">'
            f'&#x26A0;&#xFE0F; Ce previsionnel est genere automatiquement a partir des donnees '
            f'historiques du FEC. Il constitue une base de travail indicative. '
            f'L\'expert-comptable doit le revoir, l\'ajuster aux specificites du dossier '
            f'et le valider avant toute utilisation externe (financement, investissement).'
            f'</div>'
        )
    else:
        prev_content_html = ''

    html = f"""<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>FEC Audit Pro v5 — {os.path.basename(filepath)}</title>
<script src="https://cdn.jsdelivr.net/npm/chart.js@4.4.0/dist/chart.umd.min.js"></script>
<style>
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=DM+Mono:wght@400;500&family=DM+Sans:wght@300;400;500&display=swap');
:root{{
  --ink:#0d0d14;--ink2:#181825;--ink3:#22222f;--ink4:#2e2e3e;
  --line:#38384e;--muted:#6e6e90;--dim:#9898b8;--text:#ddddf0;--bright:#f0f0ff;
  --cyan:#00e5c8;--violet:#9b8afb;--red:#ff5f5f;--orange:#ff8c42;
  --yellow:#ffd166;--green:#06d6a0;--nw:252px;
}}
*{{margin:0;padding:0;box-sizing:border-box;}}
html{{scroll-behavior:smooth;}}
body{{background:var(--ink);color:var(--text);font-family:'DM Sans',sans-serif;
      display:flex;min-height:100vh;font-size:14px;}}
.sidebar{{width:var(--nw);background:var(--ink2);border-right:1px solid var(--line);
           position:fixed;top:0;left:0;bottom:0;display:flex;flex-direction:column;
           z-index:100;overflow-y:auto;}}
.slogo{{padding:20px 18px 14px;border-bottom:1px solid var(--line);}}
.slt{{font-family:'Syne',sans-serif;font-weight:800;font-size:1.05em;
      color:var(--cyan);letter-spacing:-.5px;}}
.slb{{font-family:'DM Mono',monospace;font-size:.6em;color:var(--muted);margin-top:3px;}}
.nsec{{padding:7px 0;border-bottom:1px solid var(--line);}}
.nsecl{{font-size:.6em;color:var(--muted);text-transform:uppercase;letter-spacing:1px;
         padding:5px 18px 3px;font-weight:700;}}
.ni{{display:flex;align-items:center;gap:8px;padding:7px 18px;cursor:pointer;color:var(--dim);
     font-size:.83em;border-left:2px solid transparent;transition:.15s;user-select:none;}}
.ni:hover{{color:var(--bright);background:var(--ink3);}}
.ni.active{{color:var(--cyan);background:rgba(0,229,200,.07);border-left-color:var(--cyan);font-weight:500;}}
.nico{{font-size:.95em;width:16px;text-align:center;flex-shrink:0;}}
.nbadge{{margin-left:auto;color:#fff;font-size:.62em;font-weight:700;padding:1px 6px;
          border-radius:10px;font-family:'DM Mono',monospace;flex-shrink:0;}}
.sfoot{{margin-top:auto;padding:12px 18px;font-size:.62em;color:var(--muted);
         border-top:1px solid var(--line);line-height:1.7;}}
.main{{margin-left:var(--nw);flex:1;min-width:0;}}
.topbar{{background:var(--ink2);border-bottom:1px solid var(--line);padding:13px 26px;
          display:flex;align-items:center;justify-content:space-between;
          position:sticky;top:0;z-index:50;gap:10px;flex-wrap:wrap;}}
.tbt{{font-family:'Syne',sans-serif;font-weight:700;font-size:.9em;
      white-space:nowrap;color:var(--bright);}}
.tbm{{display:flex;gap:8px;flex-wrap:wrap;}}
.chip{{background:var(--ink3);border:1px solid var(--line);padding:3px 10px;
       border-radius:20px;font-size:.7em;color:var(--dim);font-family:'DM Mono',monospace;white-space:nowrap;}}
.chip b{{color:var(--text);}}
.page{{display:none;padding:22px 26px;}}
.page.active{{display:block;animation:fadein .2s ease;}}
@keyframes fadein{{from{{opacity:.4;transform:translateY(4px);}}to{{opacity:1;transform:none;}}}}
/* SCORE HERO */
.score-hero{{display:grid;grid-template-columns:130px 1fr;gap:22px;align-items:center;
             background:linear-gradient(135deg,var(--ink2),var(--ink3));
             border:1px solid var(--line);border-radius:14px;padding:22px;margin-bottom:18px;}}
.sring{{position:relative;width:120px;height:120px;}}
.sring svg{{transform:rotate(-90deg);}}
.sring circle{{fill:none;stroke-width:10;}}
.str{{stroke:var(--ink4);}}
.stf{{stroke:{sc};stroke-dasharray:{score*2.2:.1f} 226.2;stroke-linecap:round;}}
.sctr{{position:absolute;top:50%;left:50%;transform:translate(-50%,-50%);text-align:center;}}
.snum{{font-family:'Syne',sans-serif;font-size:2em;font-weight:800;color:{sc};line-height:1;}}
.sden{{font-size:.62em;color:var(--muted);font-family:'DM Mono',monospace;}}
.sinfo h2{{font-family:'Syne',sans-serif;font-size:1.15em;font-weight:700;margin-bottom:6px;}}
.chips{{display:flex;gap:8px;flex-wrap:wrap;margin-top:10px;}}
.rchip{{padding:5px 11px;border-radius:7px;font-size:.73em;font-weight:600;
        font-family:'DM Mono',monospace;border:1px solid;}}
/* SECTEUR BADGE */
.sect-badge{{display:inline-flex;align-items:center;gap:6px;background:rgba(155,138,251,.12);
              border:1px solid rgba(155,138,251,.3);padding:4px 12px;border-radius:20px;
              font-size:.72em;color:var(--violet);font-family:'DM Mono',monospace;margin-top:6px;}}
/* KPI */
.kgrid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(135px,1fr));gap:11px;margin-bottom:18px;}}
.kpi{{background:var(--ink2);border:1px solid var(--line);border-radius:10px;padding:13px;transition:.2s;}}
.kpi:hover{{border-color:var(--violet);}}
.kl{{font-size:.66em;text-transform:uppercase;letter-spacing:.5px;color:var(--muted);margin-bottom:4px;font-weight:700;}}
.kv{{font-family:'DM Mono',monospace;font-size:1.2em;font-weight:500;}}
.ks{{font-size:.68em;color:var(--muted);margin-top:2px;}}
/* GRIDS */
.g2{{display:grid;grid-template-columns:1fr 1fr;gap:16px;margin-bottom:18px;}}
.g3{{display:grid;grid-template-columns:1fr 1fr 1fr;gap:16px;margin-bottom:18px;}}
/* BENCHMARK GAUGE */
.bench-gauge{{background:var(--ink2);border:1px solid var(--line);border-radius:12px;padding:16px;margin-bottom:18px;}}
.bench-gauge h4{{font-size:.68em;text-transform:uppercase;letter-spacing:.5px;
                  color:var(--muted);margin-bottom:14px;font-weight:700;}}
.gauge-row{{display:flex;align-items:center;gap:10px;margin-bottom:12px;}}
.gauge-label{{width:160px;font-size:.8em;color:var(--dim);flex-shrink:0;}}
.gauge-track{{flex:1;height:20px;background:var(--ink3);border-radius:10px;position:relative;overflow:visible;}}
.gauge-range{{position:absolute;height:100%;background:rgba(155,138,251,.2);
               border-radius:10px;}}
.gauge-median{{position:absolute;top:-4px;bottom:-4px;width:2px;background:var(--violet);}}
.gauge-val{{position:absolute;top:50%;width:12px;height:12px;border-radius:50%;
             transform:translate(-50%,-50%);border:2px solid;}}
.gauge-val.above{{background:var(--red);border-color:var(--red);}}
.gauge-val.inside{{background:var(--green);border-color:var(--green);}}
.gauge-val.below{{background:var(--orange);border-color:var(--orange);}}
.gauge-num{{width:60px;font-size:.82em;font-family:'DM Mono',monospace;text-align:right;}}
/* CHARTS */
.cb{{background:var(--ink2);border:1px solid var(--line);border-radius:12px;padding:15px;}}
.cb h4{{font-size:.68em;text-transform:uppercase;letter-spacing:.5px;
         color:var(--muted);margin-bottom:12px;font-weight:700;}}
.no-chart{{color:var(--muted);font-size:.8em;font-style:italic;padding:20px 0;text-align:center;}}
/* TABLES */
.sec{{background:var(--ink2);border:1px solid var(--line);border-radius:12px;margin-bottom:16px;overflow:hidden;}}
.sh{{padding:11px 18px;border-bottom:1px solid var(--line);display:flex;align-items:center;
      gap:8px;background:var(--ink3);}}
.sh h3{{font-family:'Syne',sans-serif;font-size:.78em;font-weight:700;
         text-transform:uppercase;letter-spacing:.5px;color:var(--bright);}}
.sd{{font-size:.7em;color:var(--muted);margin-left:auto;font-style:italic;}}
.sb{{padding:16px;}} .sb0{{padding:0;}}
table{{width:100%;border-collapse:collapse;font-size:.78em;}}
th{{background:var(--ink3);color:var(--muted);font-size:.65em;text-transform:uppercase;
     letter-spacing:.5px;padding:8px 12px;font-weight:700;border-bottom:1px solid var(--line);}}
td{{padding:8px 12px;border-bottom:1px solid rgba(255,255,255,.03);color:#c0c0d8;}}
tr:last-child td{{border-bottom:none;}}
tr:hover td{{background:rgba(255,255,255,.02);}}
.t-type{{font-weight:600;color:var(--text);white-space:nowrap;}}
.t-det{{font-size:.88em;color:var(--dim);}}
.t-amt{{text-align:right;font-family:'DM Mono',monospace;font-weight:600;white-space:nowrap;}}
.t-ref{{font-family:'DM Mono',monospace;font-size:.8em;color:var(--muted);}}
.empty-row{{text-align:center;color:var(--muted);padding:18px;font-style:italic;}}
.badge{{display:inline-block;padding:2px 6px;border-radius:3px;font-size:.62em;
         font-weight:700;color:#fff;text-transform:uppercase;letter-spacing:.5px;
         font-family:'DM Mono',monospace;white-space:nowrap;}}
/* MGRID */
.mgrid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(165px,1fr));gap:9px;}}
.mi{{background:var(--ink3);border-radius:7px;padding:11px 13px;}}
.ml{{font-size:.66em;color:var(--muted);margin-bottom:3px;}}
.mv{{font-family:'DM Mono',monospace;font-size:.95em;font-weight:500;}}
/* CONCENTRATION */
.cr{{display:flex;align-items:center;gap:8px;margin-bottom:8px;}}
.cn{{width:185px;font-size:.82em;color:var(--dim);overflow:hidden;
      text-overflow:ellipsis;white-space:nowrap;flex-shrink:0;}}
.cnum{{color:var(--muted);font-size:.8em;font-family:'DM Mono',monospace;}}
.cbar{{flex:1;height:7px;background:var(--ink4);border-radius:4px;overflow:hidden;}}
.cbf{{height:100%;background:linear-gradient(90deg,var(--violet),var(--cyan));
      border-radius:4px;transition:width .8s ease;}}
.cp{{width:38px;text-align:right;font-family:'DM Mono',monospace;font-size:.78em;color:var(--text);}}
.nd{{color:var(--muted);font-style:italic;font-size:.82em;padding:6px 0;}}
/* COMMENTAIRES */
.com-box{{border-radius:9px;padding:13px 16px;margin-bottom:12px;border-left:3px solid;}}
.com-title{{font-family:'Syne',sans-serif;font-weight:700;font-size:.82em;
             margin-bottom:6px;color:var(--bright);}}
.com-text{{font-size:.82em;color:var(--dim);line-height:1.65;}}
.sugg-list{{margin-top:6px;padding-left:18px;font-size:.82em;color:var(--dim);line-height:1.8;}}
.bchi{{display:inline-flex;align-items:center;gap:8px;background:var(--ink3);
        border:1px solid var(--line);padding:7px 14px;border-radius:7px;
        font-family:'DM Mono',monospace;font-size:.82em;margin-bottom:12px;}}
.bchiv{{color:var(--yellow);font-size:1.1em;font-weight:700;}}
/* NOTE SECTORIELLE */
.note-sect{{background:rgba(155,138,251,.06);border:1px solid rgba(155,138,251,.2);
             border-radius:10px;padding:14px 16px;margin-bottom:16px;font-size:.82em;
             color:var(--dim);line-height:1.7;}}
.note-sect strong{{color:var(--violet);font-weight:600;}}
/* BOUCLIER FISCAL */
.bouclier-hero{{display:grid;grid-template-columns:160px 1fr;gap:22px;align-items:center;
               background:linear-gradient(135deg,var(--ink2),var(--ink3));
               border:1px solid var(--line);border-radius:14px;padding:22px;margin-bottom:18px;}}
.brisk-ring{{position:relative;width:150px;height:150px;}}
.brisk-ring svg{{transform:rotate(-90deg);}}
.brisk-ring circle{{fill:none;stroke-width:12;}}
.brisk-bg{{stroke:var(--ink4);}}
.brisk-fill{{stroke-linecap:round;}}
.brisk-label{{position:absolute;top:50%;left:50%;transform:translate(-50%,-50%);text-align:center;}}
.brisk-num{{font-family:'Syne',sans-serif;font-size:2.3em;font-weight:800;line-height:1;}}
.brisk-sub{{font-size:.6em;color:var(--muted);font-family:'DM Mono',monospace;}}
.brisk-info h2{{font-family:'Syne',sans-serif;font-size:1.15em;font-weight:700;margin-bottom:8px;}}
.brisk-conseil{{background:rgba(255,255,255,.04);border-radius:8px;padding:10px 14px;
                font-size:.82em;color:var(--dim);line-height:1.6;margin-top:10px;
                border-left:3px solid;}}
/* Axes risque */
.axes-grid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(260px,1fr));gap:12px;margin-bottom:18px;}}
.axe-card{{background:var(--ink2);border:1px solid var(--line);border-radius:10px;padding:14px;}}
.axe-header{{display:flex;align-items:center;justify-content:space-between;margin-bottom:8px;}}
.axe-label{{font-size:.74em;font-weight:700;color:var(--text);text-transform:uppercase;letter-spacing:.4px;}}
.axe-pts{{font-family:'DM Mono',monospace;font-size:.8em;font-weight:700;}}
.axe-bar-track{{height:6px;background:var(--ink4);border-radius:3px;overflow:hidden;margin-bottom:6px;}}
.axe-bar-fill{{height:100%;border-radius:3px;transition:width .8s ease;}}
.axe-detail{{font-size:.75em;color:var(--muted);line-height:1.5;}}
/* PREVISIONNEL */
.prev-kgrid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(140px,1fr));gap:11px;margin-bottom:18px;}}
.prev-table{{width:100%;border-collapse:collapse;font-size:.82em;}}
.prev-table th{{background:var(--ink3);color:var(--muted);font-size:.65em;text-transform:uppercase;
                letter-spacing:.5px;padding:9px 14px;font-weight:700;border-bottom:1px solid var(--line);text-align:right;}}
.prev-table th:first-child{{text-align:left;}}
.prev-table td{{padding:9px 14px;border-bottom:1px solid rgba(255,255,255,.03);text-align:right;
                font-family:'DM Mono',monospace;}}
.prev-table td:first-child{{text-align:left;font-family:'DM Sans',sans-serif;color:var(--text);font-weight:500;}}
.prev-table tr.subtotal td{{color:var(--cyan);font-weight:700;background:rgba(0,229,200,.04);}}
.prev-table tr.total td{{color:var(--green);font-weight:700;font-size:1.05em;background:rgba(6,214,160,.05);
                          border-top:1px solid rgba(6,214,160,.2);}}
.prev-table tr.ded td{{color:var(--muted);}}
.var-pos{{color:var(--green)!important;}}
.var-neg{{color:var(--red)!important;}}
.hyp-box{{background:rgba(155,138,251,.06);border:1px solid rgba(155,138,251,.2);
           border-radius:10px;padding:14px 18px;margin-bottom:18px;}}
.hyp-box h4{{font-size:.7em;text-transform:uppercase;letter-spacing:.5px;color:var(--violet);
              margin-bottom:10px;font-weight:700;}}
.hyp-grid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(200px,1fr));gap:8px;}}
.hyp-item{{font-size:.8em;color:var(--dim);}}
.hyp-item b{{color:var(--text);}}
.warn-prev{{background:rgba(255,209,102,.06);border:1px solid rgba(255,209,102,.2);
             border-radius:8px;padding:10px 14px;font-size:.75em;color:var(--muted);
             line-height:1.6;margin-top:16px;font-style:italic;}}
@media(max-width:860px){{
  .sidebar{{display:none;}}.main{{margin-left:0;}}
  .g2,.g3{{grid-template-columns:1fr;}}
  .score-hero,.bouclier-hero{{grid-template-columns:1fr;}}
  .page{{padding:14px;}}
}}
</style></head>
<body>

<nav class="sidebar">
  <div class="slogo">
    <div class="slt">FEC AUDIT PRO</div>
    <div class="slb">v5.0 — 15 controles + IA</div>
  </div>
  <div class="nsec">
    <div class="nsecl">Vue generale</div>
    <div class="ni active" onclick="sp('dashboard',this)"><span class="nico">◈</span>Tableau de bord{nbadge(anomalies,'red') if nc>0 else ''}</div>
    <div class="ni" onclick="sp('toutes',this)"><span class="nico">⚠</span>Toutes les anomalies{nbadge(anomalies,'orange') if anomalies else ''}</div>
    <div class="ni" onclick="sp('benchmark',this)"><span class="nico">◎</span>Benchmark sectoriel{nbadge(cats['benchmark'],'orange') if cats['benchmark'] else ''}</div>
    <div class="ni" onclick="sp('bouclier',this)"><span class="nico">🛡</span>Bouclier Fiscal</div>
    <div class="ni" onclick="sp('previsionnel',this)"><span class="nico">📐</span>Previsionnel N+1</div>
  </div>
  <div class="nsec">
    <div class="nsecl">Statistiques</div>
    <div class="ni" onclick="sp('benford',this)"><span class="nico">📊</span>Loi de Benford{nbadge(cats['benford'],'orange') if cats['benford'] else ''}</div>
    <div class="ni" onclick="sp('marges',this)"><span class="nico">📈</span>Marges et TVA{nbadge(cats['marges'],'orange') if cats['marges'] else ''}</div>
    <div class="ni" onclick="sp('saisonnalite',this)"><span class="nico">📅</span>Saisonnalite CA{nbadge(cats['saisonnalite'],'yellow') if cats['saisonnalite'] else ''}</div>
    <div class="ni" onclick="sp('concentration',this)"><span class="nico">🎯</span>Concentration{nbadge(cats['concentration'],'orange') if cats['concentration'] else ''}</div>
  </div>
  <div class="nsec">
    <div class="nsecl">Controles ecritures</div>
    <div class="ni" onclick="sp('equilibre',this)"><span class="nico">⚖</span>Equilibre D/C{nbadge(cats['equilibre'],'red') if cats['equilibre'] else ''}</div>
    <div class="ni" onclick="sp('soldes',this)"><span class="nico">🏦</span>Soldes anormaux{nbadge(cats['soldes'],'orange') if cats['soldes'] else ''}</div>
    <div class="ni" onclick="sp('doublons',this)"><span class="nico">⊞</span>Doublons{nbadge(cats['doublons'],'orange') if cats['doublons'] else ''}</div>
    <div class="ni" onclick="sp('weekend',this)"><span class="nico">📆</span>Week-end{nbadge(cats['weekend'],'yellow') if cats['weekend'] else ''}</div>
    <div class="ni" onclick="sp('cloture',this)"><span class="nico">🔚</span>Fin d'exercice{nbadge(cats['cloture'],'orange') if cats['cloture'] else ''}</div>
    <div class="ni" onclick="sp('montants',this)"><span class="nico">💰</span>Montants suspects{nbadge(cats['ronds'],'yellow') if cats['ronds'] else ''}</div>
  </div>
  <div class="sfoot">
    Secteur : {bench.get('secteur_label','N/A')}<br>
    Genere le {datetime.now().strftime('%d/%m/%Y %H:%M')}<br>
    {len(df):,} ecritures analysees<br>
    {os.path.basename(filepath)}
  </div>
</nav>

<div class="main">
  <div class="topbar">
    <div class="tbt">📂 {os.path.basename(filepath)}</div>
    <div class="tbm">
      <span class="chip">Secteur : <b style="color:var(--violet)">{bench.get('secteur_label','N/A')}</b></span>
      <span class="chip">Periode : <b>{fe.get('date_cloture','N/A')}</b></span>
      <span class="chip">Score risque : <b style="color:{sc}">{score}/100 — {sl}</b></span>
      <span class="chip">Bouclier DGFiP : <b style="color:{b_color}">{b_score}/100 — {b_niveau}</b></span>
      <span class="chip">Anomalies : <b>{len(anomalies)}</b></span>
    </div>
  </div>

  <!-- DASHBOARD -->
  <div class="page active" id="pg-dashboard">
    <div class="score-hero">
      <div class="sring">
        <svg width="120" height="120" viewBox="0 0 120 120">
          <circle class="str" cx="60" cy="60" r="36"/>
          <circle class="stf" cx="60" cy="60" r="36"/>
        </svg>
        <div class="sctr"><div class="snum">{score}</div><div class="sden">/100</div></div>
      </div>
      <div class="sinfo">
        <h2>Risque : <span style="color:{sc}">{sl}</span></h2>
        <p style="color:var(--dim);font-size:.83em">{len(anomalies)} anomalie(s) sur {len(df):,} ecritures — 15 controles</p>
        <div class="sect-badge">◎ {bench.get('secteur_label','Secteur inconnu')}</div>
        <div class="chips" style="margin-top:10px">
          <span class="rchip" style="background:rgba(239,68,68,.1);border-color:#ef4444;color:#ef4444">{nc} Critique{'s' if nc>1 else ''}</span>
          <span class="rchip" style="background:rgba(249,115,22,.1);border-color:#f97316;color:#f97316">{na} Alerte{'s' if na>1 else ''}</span>
          <span class="rchip" style="background:rgba(234,179,8,.1);border-color:#eab308;color:#eab308">{nat} Attention</span>
          <span class="rchip" style="background:rgba(34,197,94,.1);border-color:#22c55e;color:#22c55e">{ni} Info</span>
        </div>
      </div>
    </div>
    {gcom_html}
    {sugg_html}
    <div class="kgrid">
      <div class="kpi"><div class="kl">Chiffre d'affaires</div><div class="kv" style="color:var(--cyan)">{mg.get('ca',0):,.0f}</div><div class="ks">EUR</div></div>
      <div class="kpi"><div class="kl">Marge brute</div><div class="kv" style="color:{'var(--green)' if mg.get('mb',0)>=0 else 'var(--red)'}">{mg.get('mb',0):,.0f}</div><div class="ks">Taux {mg.get('tm',0):.1f}% | Sect. {tm_med:.0f}%</div></div>
      <div class="kpi"><div class="kl">Resultat net</div><div class="kv" style="color:{'var(--green)' if mg.get('res',0)>=0 else 'var(--red)'}">{mg.get('res',0):,.0f}</div><div class="ks">{mg.get('tr',0):.1f}% du CA</div></div>
      <div class="kpi"><div class="kl">Benford Chi2</div><div class="kv" style="color:{'var(--red)' if ben.get('chi2',0)>15.5 else 'var(--green)'}">{ben.get('chi2',0):.1f}</div><div class="ks">Seuil 15.5</div></div>
      <div class="kpi"><div class="kl">Ecr. week-end</div><div class="kv" style="color:{'var(--yellow)' if wd.get('pct',0)>5 else 'var(--text)'}">{wd.get('nb',0)}</div><div class="ks">{wd.get('pct',0):.1f}% total</div></div>
      <div class="kpi"><div class="kl">TVA collectee</div><div class="kv">{mg.get('tva_c',0):,.0f}</div><div class="ks">Taux {mg.get('tx_tva',0):.1f}%</div></div>
      <div class="kpi"><div class="kl">Fin exercice 7j</div><div class="kv" style="color:{'var(--red)' if fe.get('p7',0)>25 else 'var(--text)'}">{fe.get('p7',0):.0f}%</div><div class="ks">du volume</div></div>
      <div class="kpi"><div class="kl">Doublons</div><div class="kv" style="color:{'var(--red)' if stats.get('doublons',{}).get('nb',0)>0 else 'var(--green)'}">{stats.get('doublons',{}).get('nb',0)}</div><div class="ks">lignes</div></div>
      <div class="kpi"><div class="kl">IS comptabilise</div><div class="kv">{fisc.get('is_d',0):,.0f}</div><div class="ks">Taux {fisc.get('tx_is',0):.1f}%</div></div>
      <div class="kpi" style="cursor:pointer" onclick="sp('bouclier',document.querySelector('[onclick*=bouclier]'))"><div class="kl">Bouclier Fiscal &#x1F6E1;</div><div class="kv" style="color:{b_color}">{b_score}/100</div><div class="ks">{b_niveau}</div></div>
      <div class="kpi" style="cursor:pointer" onclick="sp('previsionnel',document.querySelector('[onclick*=previsionnel]'))"><div class="kl">CA previsionnel N+1 &#x1F4D0;</div><div class="kv" style="color:var(--cyan)">{p_ca_n1:,.0f}</div><div class="ks">EUR (+{p_croiss*100:.0f}%)</div></div>
    </div>
    <div class="g2">
      {chart_box('cMois','Volume ecritures par mois',has_mois,'Donnees insuffisantes (moins de 2 mois)',180)}
      {chart_box('cGravite','Repartition des anomalies',has_gravite,'Aucune anomalie a representer',180) if has_gravite else '<div class="cb"><h4>Repartition des anomalies</h4><div class="no-chart">Aucune anomalie detectee</div></div>'}
    </div>
    <div class="g2">
      {chart_box('cClasses','Soldes D/C par classe de compte (k EUR)',has_classes,'Comptes insuffisants',180)}
      {chart_box('cJours','Ecritures par jour de semaine',has_jours,'Dates manquantes dans le FEC',180)}
    </div>
  </div>

  <!-- TOUTES ANOMALIES -->
  <div class="page" id="pg-toutes">
    <div class="sec"><div class="sh"><span>⚠️</span><h3>Toutes les anomalies ({len(anomalies)})</h3><span class="sd">Triees par gravite</span></div>
    <div class="sb0"><table><thead><tr><th>Gravite</th><th>Type</th><th>Detail</th><th>Montant</th><th>Ref.</th></tr></thead><tbody>{rows(anomalies)}</tbody></table></div></div>
    {sugg_html}
  </div>

  <!-- BENCHMARK SECTORIEL (v4/v5) -->
  <div class="page" id="pg-benchmark">
    {comment_html('benchmark')}
    <div class="note-sect">
      <strong>Secteur : {bench.get('secteur_label','N/A')}</strong><br>
      {bench.get('note_sectorielle','')}
    </div>
    <div class="g2">
      {chart_box('cBenchTM','Marge brute : entreprise vs secteur (%)',has_bench,'CA non disponible',160)}
      {chart_box('cBenchCP','Charges personnel/CA : entreprise vs secteur (%)',has_bench,'Donnees insuffisantes',160)}
    </div>
    <div class="sec"><div class="sh"><span>◎</span><h3>Indicateurs comparatifs</h3></div>
    <div class="sb"><div class="mgrid">{bench_items}</div></div></div>
    <div class="sec"><div class="sh"><span>⚠️</span><h3>Anomalies benchmark ({len(cats['benchmark'])})</h3></div>
    <div class="sb0"><table><thead><tr><th>Gravite</th><th>Type</th><th>Detail</th><th>Montant</th><th>Ref.</th></tr></thead><tbody>{rows(cats['benchmark'])}</tbody></table></div></div>
  </div>

  <!-- BENFORD -->
  <div class="page" id="pg-benford">
    {comment_html('benford')}
    <div class="bchi"><span>Chi2 :</span><span class="bchiv">{ben.get('chi2',0):.2f}</span><span style="color:var(--muted)">/ seuil 15.5</span><span style="color:{'var(--red)' if ben.get('chi2',0)>15.5 else 'var(--green)'}">{'SUSPECT' if ben.get('chi2',0)>15.5 else 'Normal'}</span></div>
    <div class="g2">
      {chart_box('cBenford','Distribution premiers chiffres — Observe vs Attendu (%)',has_benford,'Benford non applicable (moins de 100 montants)',220)}
      {chart_box('cBenfordE','Ecarts en points %',has_benford,'Benford non applicable',220)}
    </div>
    <div class="sec"><div class="sh"><span>🔎</span><h3>Anomalies Benford ({len(cats['benford'])})</h3></div>
    <div class="sb0"><table><thead><tr><th>Gravite</th><th>Type</th><th>Detail</th><th>Montant</th><th>Ref.</th></tr></thead><tbody>{rows(cats['benford'])}</tbody></table></div></div>
  </div>

  <!-- MARGES -->
  <div class="page" id="pg-marges">
    {comment_html('marges')}
    <div class="sec"><div class="sh"><span>📈</span><h3>Marges, resultats et TVA</h3></div>
    <div class="sb"><div class="mgrid">{marge_items}</div></div></div>
    <div class="sec"><div class="sh"><span>⚠️</span><h3>Anomalies ({len(cats['marges'])})</h3></div>
    <div class="sb0"><table><thead><tr><th>Gravite</th><th>Type</th><th>Detail</th><th>Montant</th><th>Ref.</th></tr></thead><tbody>{rows(cats['marges'])}</tbody></table></div></div>
  </div>

  <!-- SAISONNALITE -->
  <div class="page" id="pg-saisonnalite">
    {comment_html('saisonnalite')}
    <div class="g2">
      {chart_box('cCA','CA mensuel (k EUR)',has_ca,'Aucune donnee de ventes mensuelle',200)}
      {chart_box('cMois2','Nombre d\'ecritures par mois',has_mois,'Donnees insuffisantes',200)}
    </div>
    <div class="sec"><div class="sh"><span>⚠️</span><h3>Anomalies saisonnalite ({len(cats['saisonnalite'])})</h3></div>
    <div class="sb0"><table><thead><tr><th>Gravite</th><th>Type</th><th>Detail</th><th>Montant</th><th>Ref.</th></tr></thead><tbody>{rows(cats['saisonnalite'])}</tbody></table></div></div>
  </div>

  <!-- CONCENTRATION -->
  <div class="page" id="pg-concentration">
    {comment_html('concentration')}
    <div class="g2">
      <div class="cb"><h4>Top clients — % du CA</h4><div style="padding:2px 0">{top_html}</div></div>
      {chart_box('cConc','Repartition CA par tiers',has_conc,'Donnees CompAuxNum absentes',200)}
    </div>
    <div class="sec"><div class="sh"><span>⚠️</span><h3>Anomalies concentration ({len(cats['concentration'])})</h3></div>
    <div class="sb0"><table><thead><tr><th>Gravite</th><th>Type</th><th>Detail</th><th>Montant</th><th>Ref.</th></tr></thead><tbody>{rows(cats['concentration'])}</tbody></table></div></div>
  </div>

  <!-- EQUILIBRE -->
  <div class="page" id="pg-equilibre">
    {comment_html('equilibre')}
    <div class="sec"><div class="sh"><span>⚖️</span><h3>Anomalies equilibre D/C ({len(cats['equilibre'])})</h3></div>
    <div class="sb0"><table><thead><tr><th>Gravite</th><th>Type</th><th>Detail</th><th>Montant</th><th>Ref.</th></tr></thead><tbody>{rows(cats['equilibre'])}</tbody></table></div></div>
  </div>

  <!-- SOLDES -->
  <div class="page" id="pg-soldes">
    {comment_html('soldes')}
    <div class="sec"><div class="sh"><span>🏦</span><h3>Anomalies de soldes ({len(cats['soldes'])})</h3></div>
    <div class="sb0"><table><thead><tr><th>Gravite</th><th>Type</th><th>Detail</th><th>Montant</th><th>Ref.</th></tr></thead><tbody>{rows(cats['soldes'])}</tbody></table></div></div>
  </div>

  <!-- DOUBLONS -->
  <div class="page" id="pg-doublons">
    {comment_html('doublons')}
    <div class="sec"><div class="sh"><span>⊞</span><h3>Doublons et ecritures inversees ({len(cats['doublons'])})</h3></div>
    <div class="sb0"><table><thead><tr><th>Gravite</th><th>Type</th><th>Detail</th><th>Montant</th><th>Ref.</th></tr></thead><tbody>{rows(cats['doublons'])}</tbody></table></div></div>
  </div>

  <!-- WEEK-END -->
  <div class="page" id="pg-weekend">
    {comment_html('weekend')}
    <div class="g2">
      {chart_box('cJours2','Ecritures par jour de la semaine',has_jours,'Dates non disponibles',200)}
      <div class="cb" style="display:flex;flex-direction:column;justify-content:center">
        <h4>Statistiques</h4>
        <div style="margin-top:10px">
          <div class="mi" style="margin-bottom:8px"><div class="ml">Ecritures week-end</div><div class="mv">{wd.get('nb',0)}</div></div>
          <div class="mi"><div class="ml">% du total</div><div class="mv">{wd.get('pct',0):.1f}%</div></div>
        </div>
      </div>
    </div>
    <div class="sec"><div class="sh"><span>📆</span><h3>Anomalies week-end ({len(cats['weekend'])})</h3></div>
    <div class="sb0"><table><thead><tr><th>Gravite</th><th>Type</th><th>Detail</th><th>Montant</th><th>Ref.</th></tr></thead><tbody>{rows(cats['weekend'])}</tbody></table></div></div>
  </div>

  <!-- CLOTURE -->
  <div class="page" id="pg-cloture">
    {comment_html('cloture')}
    <div class="g2">
      {chart_box('cFe','Volume fin d\'exercice',has_fe,'Dates non disponibles',200)}
      <div class="cb" style="display:flex;flex-direction:column;justify-content:center">
        <h4>Indicateurs</h4>
        <div style="margin-top:10px">
          <div class="mi" style="margin-bottom:8px"><div class="ml">% volume derniers 7j</div><div class="mv" style="color:{'var(--red)' if fe.get('p7',0)>25 else 'var(--text)'}">{fe.get('p7',0):.1f}%</div></div>
          <div class="mi" style="margin-bottom:8px"><div class="ml">% volume derniers 30j</div><div class="mv">{fe.get('p30',0):.1f}%</div></div>
          <div class="mi"><div class="ml">Date de cloture</div><div class="mv">{fe.get('date_cloture','N/A')}</div></div>
        </div>
      </div>
    </div>
    <div class="sec"><div class="sh"><span>🔚</span><h3>Anomalies fin d'exercice ({len(cats['cloture'])})</h3></div>
    <div class="sb0"><table><thead><tr><th>Gravite</th><th>Type</th><th>Detail</th><th>Montant</th><th>Ref.</th></tr></thead><tbody>{rows(cats['cloture'])}</tbody></table></div></div>
  </div>

  <!-- MONTANTS -->
  <div class="page" id="pg-montants">
    {comment_html('montants')}
    <div class="sec"><div class="sh"><span>💰</span><h3>Montants suspects ({len(cats['ronds'])})</h3><span class="sd">Multiples de 1000 EUR : {mr.get('p1000',0):.1f}%</span></div>
    <div class="sb0"><table><thead><tr><th>Gravite</th><th>Type</th><th>Detail</th><th>Montant</th><th>Ref.</th></tr></thead><tbody>{rows(cats['ronds'])}</tbody></table></div></div>
  </div>  <!-- BOUCLIER FISCAL (Module 2 v5) -->
  <div class="page" id="pg-bouclier">
    <div class="bouclier-hero">
      <div class="brisk-ring">
        <svg width="150" height="150" viewBox="0 0 150 150">
          <circle class="brisk-bg" cx="75" cy="75" r="54"/>
          <circle class="brisk-fill" cx="75" cy="75" r="54"
            style="stroke:{b_color};stroke-dasharray:{b_dash:.1f} 339.3;stroke-width:12"/>
        </svg>
        <div class="brisk-label">
          <div class="brisk-num" style="color:{b_color}">{b_score}</div>
          <div class="brisk-sub">/100</div>
        </div>
      </div>
      <div class="brisk-info">
        <h2>Probabilite de controle : <span style="color:{b_color}">{b_niveau}</span></h2>
        <p style="color:var(--dim);font-size:.83em">{bouclier.get('nb_axes_risque',0)} axe(s) de risque actifs sur {len(b_details)}</p>
        <div class="brisk-conseil" style="border-color:{b_color}">{b_conseil}</div>
      </div>
    </div>

    <div class="axes-grid">
      {axes_html if axes_html else '<p style="color:var(--muted);font-style:italic">Donnees insuffisantes pour calculer le score.</p>'}
    </div>

    <div class="sec">
      <div class="sh"><span>🛡</span><h3>Methodologie du score</h3></div>
      <div class="sb" style="font-size:.82em;color:var(--dim);line-height:1.8">
        Le <strong style="color:var(--text)">Score de Probabilite de Controle DGFiP</strong>
        est calcule sur 8 axes de risque independants, inspires des criteres
        de ciblage de la Direction Generale des Finances Publiques (CFCI) :<br><br>
        <strong style="color:var(--violet)">Loi de Benford (25 pts)</strong> — Detecte les manipulations statistiques des montants
        (saisies manuelles, forfaits, montants inventes).<br>
        <strong style="color:var(--violet)">Coherence TVA (20 pts)</strong> — Analyse le ratio TVA deductible / collectee
        et la coherence avec le taux sectoriel attendu.<br>
        <strong style="color:var(--violet)">Concentration fin d'exercice (15 pts)</strong> — Signal d'habillage de bilan :
        concentration anormale d'ecritures dans les derniers jours de l'exercice.<br>
        <strong style="color:var(--violet)">Marge vs secteur (12 pts)</strong> — Une marge tres eloignee de la mediane sectorielle
        peut indiquer des recettes dissimulees ou des charges surestimees.<br>
        <strong style="color:var(--violet)">Impot sur les Societes (10 pts)</strong> — Absence d'IS malgre un resultat positif
        ou taux IS anormalement faible ou eleve.<br>
        <strong style="color:var(--violet)">Montants aberrants (10 pts)</strong> — Montants a plus de 3 ecarts-types
        de la moyenne, superieurs a 5 000 EUR.<br>
        <strong style="color:var(--violet)">Charges salariales (8 pts)</strong> — Ratio cotisations / salaires bruts
        tres eloigne de la norme 42-48%.<br>
        <strong style="color:var(--violet)">Integrite FEC (5 pts)</strong> — Doublons, importations multiples ou
        autres signes d'alteration du fichier.<br><br>
        <em>Ce score est un indicateur d'aide a la decision. Il ne se substitue pas
        au jugement professionnel de l'expert-comptable.</em>
      </div>
    </div>
  </div>

  <!-- PREVISIONNEL N+1 (Module 3 v5) -->
  <div class="page" id="pg-previsionnel">
    {prev_no_data_html}
    {prev_content_html}
  </div>

</div>

<script>
Chart.defaults.color='#6e6e90';
Chart.defaults.font.family="'DM Mono',monospace";
Chart.defaults.font.size=11;
const C={{cyan:'#00e5c8',violet:'#9b8afb',red:'#ff5f5f',orange:'#ff8c42',yellow:'#ffd166',green:'#06d6a0',line:'#38384e'}};

function sp(id,el){{
  document.querySelectorAll('.page').forEach(p=>p.classList.remove('active'));
  document.querySelectorAll('.ni').forEach(n=>n.classList.remove('active'));
  document.getElementById('pg-'+id).classList.add('active');
  el.classList.add('active');
}}

function mkBar(id,labels,data,color,opts={{}}){{
  const el=document.getElementById(id); if(!el||!labels.length) return;
  new Chart(el,{{type:'bar',data:{{labels,datasets:[{{label:'',data,backgroundColor:color+'99',borderColor:color,borderWidth:1,borderRadius:3}}]}},
    options:{{responsive:true,plugins:{{legend:{{display:false}}}},scales:{{x:{{grid:{{color:C.line}},ticks:{{maxRotation:45}}}},y:{{grid:{{color:C.line}}}}}}, ...opts}}}});
}}

// Volume mensuel
const lm={lm},dm={dm};
if(lm.length){{mkBar('cMois',lm,dm,C.violet);mkBar('cMois2',lm,dm,C.cyan);}}

// Gravites
if({1 if has_gravite else 0}){{
  const eg=document.getElementById('cGravite');
  if(eg)new Chart(eg,{{type:'doughnut',data:{{labels:['Critique','Alerte','Attention','Info'],datasets:[{{data:[{nc},{na},{nat},{ni}],backgroundColor:['rgba(255,95,95,.8)','rgba(255,140,66,.8)','rgba(255,209,102,.8)','rgba(6,214,160,.8)'],borderColor:'#181825',borderWidth:2}}]}},options:{{responsive:true,plugins:{{legend:{{position:'right',labels:{{padding:10}}}}}}}}}});
}}

// Classes
if({1 if has_classes else 0}){{
  const ecl=document.getElementById('cClasses');
  if(ecl)new Chart(ecl,{{type:'bar',data:{{labels:{cl_labels},datasets:[{{label:'Debit(k)',data:{cl_d},backgroundColor:'rgba(255,95,95,.6)',borderColor:C.red,borderWidth:1,borderRadius:2}},{{label:'Credit(k)',data:{cl_c},backgroundColor:'rgba(6,214,160,.6)',borderColor:C.green,borderWidth:1,borderRadius:2}}]}},options:{{responsive:true,plugins:{{legend:{{position:'top'}}}},scales:{{x:{{grid:{{color:C.line}}}},y:{{grid:{{color:C.line}}}}}}}}}});
}}

// Jours
const lj={ljours},dj={djours};
if(lj.length){{
  ['cJours','cJours2'].forEach(id=>{{
    const el=document.getElementById(id); if(!el) return;
    new Chart(el,{{type:'bar',data:{{labels:lj,datasets:[{{label:'',data:dj,backgroundColor:lj.map((_,i)=>i>=5?'rgba(255,140,66,.65)':'rgba(0,229,200,.45)'),borderColor:lj.map((_,i)=>i>=5?C.orange:C.cyan),borderWidth:1,borderRadius:3}}]}},options:{{responsive:true,plugins:{{legend:{{display:false}}}},scales:{{x:{{grid:{{color:C.line}}}},y:{{grid:{{color:C.line}}}}}}}}}});
  }});
}}

// CA mensuel
if({1 if has_ca else 0}){{
  const eca=document.getElementById('cCA');
  if(eca)new Chart(eca,{{type:'line',data:{{labels:{lca},datasets:[{{label:'CA(k)',data:{dca},borderColor:C.cyan,backgroundColor:'rgba(0,229,200,.1)',fill:true,tension:.3,pointRadius:3,pointBackgroundColor:C.cyan}}]}},options:{{responsive:true,plugins:{{legend:{{display:false}}}},scales:{{x:{{grid:{{color:C.line}}}},y:{{grid:{{color:C.line}}}}}}}}}});
}}

// Benford
if({1 if has_benford else 0}){{
  const bl={b_labels},bo={b_obs},ba={b_att};
  const eb=document.getElementById('cBenford');
  if(eb)new Chart(eb,{{type:'bar',data:{{labels:bl,datasets:[{{label:'Observe %',data:bo,backgroundColor:'rgba(155,138,251,.6)',borderColor:C.violet,borderWidth:1,borderRadius:2}},{{label:'Attendu %',data:ba,type:'line',borderColor:C.cyan,backgroundColor:'transparent',tension:.3,pointRadius:3,borderWidth:2}}]}},options:{{responsive:true,plugins:{{legend:{{position:'top'}}}},scales:{{x:{{grid:{{color:C.line}}}},y:{{grid:{{color:C.line}},title:{{display:true,text:'%'}}}}}}}}}});
  const eb2=document.getElementById('cBenfordE');
  if(eb2){{const ec=bo.map((o,i)=>parseFloat((o-ba[i]).toFixed(2)));new Chart(eb2,{{type:'bar',data:{{labels:bl,datasets:[{{data:ec,backgroundColor:ec.map(e=>Math.abs(e)>5?'rgba(255,95,95,.7)':'rgba(6,214,160,.5)'),borderColor:ec.map(e=>Math.abs(e)>5?C.red:C.green),borderWidth:1,borderRadius:2}}]}},options:{{responsive:true,plugins:{{legend:{{display:false}}}},scales:{{x:{{grid:{{color:C.line}}}},y:{{grid:{{color:C.line}},title:{{display:true,text:'Ecart pts %'}}}}}}}}}});}}
}}

// Concentration
if({1 if has_conc else 0}){{
  const ecc=document.getElementById('cConc');
  const tcd={tc_json};
  if(ecc&&tcd.length)new Chart(ecc,{{type:'pie',data:{{labels:tcd.map(c=>c.lib||'?'),datasets:[{{data:tcd.map(c=>c.pct),backgroundColor:['rgba(0,229,200,.8)','rgba(155,138,251,.8)','rgba(255,140,66,.8)','rgba(255,209,102,.8)','rgba(255,95,95,.8)'],borderColor:'#181825',borderWidth:2}}]}},options:{{responsive:true,plugins:{{legend:{{position:'right',labels:{{padding:10,font:{{size:10}}}}}}}}}}}});
}}

// Fin exercice
if({1 if has_fe else 0}){{
  const efe=document.getElementById('cFe');
  if(efe)new Chart(efe,{{type:'doughnut',data:{{labels:['Derniers 7j','Reste annee'],datasets:[{{data:[{p7:.1f},{p_rest:.1f}],backgroundColor:['rgba(255,95,95,.8)','rgba(155,138,251,.4)'],borderColor:'#181825',borderWidth:2}}]}},options:{{responsive:true,plugins:{{legend:{{position:'right'}}}}}}}});
}}

// Benchmark sectoriel — graphes comparatifs
if({1 if has_bench else 0}){{
  const ebtm=document.getElementById('cBenchTM');
  if(ebtm)new Chart(ebtm,{{type:'bar',
    data:{{labels:{bench_chart_labels},datasets:[{{data:{bench_chart_data},backgroundColor:['rgba(0,229,200,.7)','rgba(155,138,251,.5)','rgba(255,209,102,.4)','rgba(255,209,102,.4)'],borderColor:[C.cyan,C.violet,C.yellow,C.yellow],borderWidth:1,borderRadius:3}}]}},
    options:{{responsive:true,plugins:{{legend:{{display:false}}}},scales:{{x:{{grid:{{color:C.line}}}},y:{{grid:{{color:C.line}},title:{{display:true,text:'%'}}}}}}}}}});
  const ebcp=document.getElementById('cBenchCP');
  if(ebcp)new Chart(ebcp,{{type:'bar',
    data:{{labels:{cp_chart_labels},datasets:[{{data:{cp_chart_data},backgroundColor:['rgba(255,140,66,.7)','rgba(155,138,251,.5)'],borderColor:[C.orange,C.violet],borderWidth:1,borderRadius:3}}]}},
    options:{{responsive:true,plugins:{{legend:{{display:false}}}},scales:{{x:{{grid:{{color:C.line}}}},y:{{grid:{{color:C.line}},title:{{display:true,text:'%'}}}}}}}}}});
}}

// Previsionnel CA mensuel N vs N+1
if({1 if has_prev_ca else 0}){{
  const epca=document.getElementById('cPrevCA');
  if(epca)new Chart(epca,{{type:'bar',
    data:{{labels:{lca_n1},datasets:[
      {{label:'{p_annee_n} (reel)',data:{dca_n_prev},backgroundColor:'rgba(155,138,251,.5)',borderColor:C.violet,borderWidth:1,borderRadius:3}},
      {{label:'{p_annee_n1} (prev.)',data:{dca_n1},backgroundColor:'rgba(0,229,200,.5)',borderColor:C.cyan,borderWidth:1,borderRadius:3}}
    ]}},
    options:{{responsive:true,plugins:{{legend:{{position:'top'}}}},
             scales:{{x:{{grid:{{color:C.line}}}},y:{{grid:{{color:C.line}},title:{{display:true,text:'k EUR'}}}}}}}}
  }});
}}
</script>
</body></html>"""
    return html

# ══════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════

# ══════════════════════════════════════════════════════════════
# FIN DES MODULES V5
# ══════════════════════════════════════════════════════════════

# ══════════════════════════════════════════════════════════════
# MODULE 1 — GENERATEUR DE RAPPORT DE MISSION IA
# ══════════════════════════════════════════════════════════════

def _construire_prompt_rapport(stats, anomalies, secteur_code, nom_fichier):
    bench    = stats.get('benchmark', {})
    mg       = stats.get('marges', {})
    sal      = stats.get('charges_salariales', {})
    fisc     = stats.get('scoring_fiscal', {})
    bouclier = stats.get('bouclier_fiscal', {})

    lignes_anom = []
    for a in anomalies:
        if a['gravite'] in ('CRITIQUE', 'ALERTE'):
            lignes_anom.append(f"- [{a['gravite']}] {a['type']} : {a['detail']}")
    anom_txt = "\n".join(lignes_anom[:12]) if lignes_anom else "Aucune anomalie majeure."

    ca      = mg.get('ca', 0)
    tm      = mg.get('tm', 0)
    tm_med  = bench.get('tm_mediane', bench.get('tm_med', 0))
    tr      = mg.get('tr', 0)
    cp      = mg.get('cp', 0)
    res     = mg.get('res', 0)
    nb_sal  = sal.get('nb_sal_estime', 0)
    tx_is   = fisc.get('tx_is', 0)
    b_score = bouclier.get('score_global', 0)
    b_niv   = bouclier.get('niveau_risque', 'N/A')
    sect_lb = BENCHMARKS.get(secteur_code, {}).get('label', secteur_code)

    return f"""Tu es un expert-comptable senior francais.
Redige un rapport de mission d'audit comptable professionnel en francais, destine au client (dirigeant de PME).

CONTEXTE :
- Fichier : {nom_fichier}
- Secteur : {sect_lb}
- CA : {ca:,.0f} EUR
- Marge brute : {tm:.1f}% (mediane sectorielle {tm_med:.0f}%)
- Resultat : {res:,.0f} EUR ({tr:.1f}% du CA)
- Charges personnel : {cp:,.0f} EUR (effectif ~{nb_sal})
- Taux IS apparent : {tx_is:.1f}%
- Score Bouclier Fiscal DGFiP : {b_score}/100 ({b_niv})

ANOMALIES DETECTEES :
{anom_txt}

STRUCTURE ATTENDUE (rapport de 3 pages) :
1. Introduction avec formule de politesse professionnelle.
2. "Synthese de notre mission" (2-3 §) : contexte, objectifs, methode.
3. "Analyse des performances financieres" (3-4 §) : CA, marges, resultat, charges, comparaison sectorielle chiffree.
4. "Points d'attention et risques identifies" : chaque anomalie expliquee de facon pedagogique pour un non-comptable, consequences potentielles, actions correctrices.
5. "Recommandations et plan d'action" : 4 a 6 recommandations concretes priorisees.
6. Conclusion et formule de cloture professionnelle.

STYLE : Professionnel mais accessible. Pas de tutoiement. Chiffre les ecarts. Direct sur les risques sans alarmer inutilement."""


def generer_rapport_mission_ia(filepath, stats, anomalies, secteur_code,
                                api_key=None, output_docx=None):
    """
    MODULE 1 — Rapport de Mission IA.
    Envoie les resultats a l'API Claude et genere un Word professionnel.
    Necessite : pip install anthropic python-docx
    """
    if not ANTHROPIC_AVAILABLE:
        msg = ("MODULE 1 indisponible : installez le package 'anthropic'\n"
               "  pip install anthropic")
        print(f"\n  {msg}"); return None, msg

    key = api_key or os.environ.get('ANTHROPIC_API_KEY', '')
    if not key:
        msg = ("MODULE 1 : cle API manquante.\n"
               "  Definissez ANTHROPIC_API_KEY ou passez --api-key sk-ant-...")
        print(f"\n  {msg}"); return None, msg

    nom = os.path.basename(filepath)
    print(f"\n  [MODULE 1] Rapport de mission IA — envoi a Claude...")
    try:
        client  = _anthropic_lib.Anthropic(api_key=key)
        prompt  = _construire_prompt_rapport(stats, anomalies, secteur_code, nom)
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=3000,
            messages=[{"role": "user", "content": prompt}]
        )
        texte = message.content[0].text
        print(f"  Rapport genere ({len(texte)} caracteres).")
    except Exception as e:
        msg = f"MODULE 1 erreur API : {e}"
        print(f"  {msg}"); return None, msg

    if output_docx is None:
        output_docx = os.path.splitext(filepath)[0] + "_rapport_mission.docx"

    if DOCX_AVAILABLE:
        try:
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            doc = _DocxDocument()
            for sec in doc.sections:
                sec.top_margin = sec.bottom_margin = Cm(2.5)
                sec.left_margin = Cm(3.0); sec.right_margin = Cm(2.5)
            t = doc.add_heading(
                "Rapport de Mission — Analyse du Fichier des Ecritures Comptables", 0)
            t.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
            p = doc.add_paragraph()
            p.add_run(
                f"Date : {datetime.now().strftime('%d/%m/%Y')}    "
                f"Fichier : {nom}    "
                f"Secteur : {BENCHMARKS.get(secteur_code, {}).get('label', secteur_code)}"
            ).font.size = Pt(9)
            doc.add_paragraph()
            for ligne in texte.split("\n"):
                l = ligne.strip()
                if not l:
                    doc.add_paragraph(); continue
                if re.match(r'^(Section\s+\d+|[1-9][\.)\s])', l):
                    h = doc.add_heading(l, level=1)
                    h.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
                elif l.startswith(("- ", "• ")):
                    doc.add_paragraph(l.lstrip("-• "), style="List Bullet")
                else:
                    doc.add_paragraph(l)
            doc.add_paragraph()
            pf = doc.add_paragraph(
                "Document genere par FEC Audit Pro v5.0 — "
                "A relire et valider par l'expert-comptable avant envoi.")
            pf.runs[0].font.size = Pt(8)
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.save(output_docx)
            print(f"  Word sauvegarde : {output_docx}")
        except Exception as e:
            txt_path = os.path.splitext(output_docx)[0] + ".txt"
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(f"RAPPORT DE MISSION — {nom}\n{'='*70}\n\n{texte}")
            print(f"  (docx echoue: {e}) TXT : {txt_path}")
            output_docx = txt_path
    else:
        txt_path = os.path.splitext(output_docx)[0] + ".txt"
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(f"RAPPORT DE MISSION — {nom}\n{'='*70}\n\n{texte}")
        print(f"  (python-docx absent) TXT : {txt_path}")
        output_docx = txt_path

    return output_docx, texte


# ══════════════════════════════════════════════════════════════
# MODULE 2 — BOUCLIER FISCAL
# ══════════════════════════════════════════════════════════════

GRILLE_RISQUE_FISCAL = [
    (80, 'TRES ELEVE',  'Pre-audit fiscal recommande en urgence.'),
    (60, 'ELEVE',       'Revue contradictoire des points identifies recommandee.'),
    (40, 'MODERE',      'Renforcer la documentation des points sensibles.'),
    (20, 'FAIBLE',      'Maintenir la rigueur de saisie et de justification.'),
    ( 0, 'TRES FAIBLE', 'Aucun signal majeur. Dossier conforme aux normes DGFiP.'),
]

def calculer_bouclier_fiscal(df, stats, anomalies, secteur_code):
    """
    MODULE 2 — Score de Probabilite de Controle DGFiP (0-100).
    8 axes : Benford, TVA, fin exercice, aberrants, marge/secteur,
             IS, charges salariales, integrite FEC.
    """
    bench = stats.get('benchmark', {}); ben  = stats.get('benford', {})
    fe    = stats.get('fin_exercice', {}); mg  = stats.get('marges', {})
    fisc  = stats.get('scoring_fiscal', {}); sal = stats.get('charges_salariales', {})
    ab    = stats.get('aberrants', {}); dup  = stats.get('doublons', {})
    details = {}

    # 1. BENFORD (25 pts)
    chi2 = ben.get('chi2', 0); tot = ben.get('total', 0)
    if tot >= 100:
        if chi2 > 25:   pts, lbl = 25, f"Distribution tres anormale (Chi2={chi2:.1f})"
        elif chi2 > 15.5: pts, lbl = 15, f"Distribution suspecte (Chi2={chi2:.1f})"
        elif chi2 > 10:  pts, lbl =  6, f"Legere deviation (Chi2={chi2:.1f})"
        else:            pts, lbl =  0, f"Distribution conforme (Chi2={chi2:.1f})"
    else:                pts, lbl =  0, "Non applicable (< 100 montants)"
    details['benford'] = {'points': pts, 'max': 25, 'label': 'Loi de Benford', 'detail': lbl}

    # 2. TVA (20 pts)
    tva_c = mg.get('tva_c', 0); tva_d = mg.get('tva_d', 0)
    tx_tva = mg.get('tx_tva', 0)
    btva = BENCHMARKS.get(secteur_code, {}).get('tx_tva', 20)
    pts = 0; lbl = "TVA coherente"
    if tva_c > 0:
        r = tva_d / tva_c
        if r > 2:         pts, lbl = 20, f"Credit TVA structurel (deductible={r:.1f}x collectee)"
        elif r > 1.5:     pts, lbl = 12, f"TVA deductible elevee (x{r:.1f})"
        elif btva == 0 and tx_tva > 1: pts, lbl = 15, "TVA inattendue (secteur exonere)"
        elif btva > 0 and abs(tx_tva - btva) > 8:
            pts, lbl = 10, f"Taux TVA ({tx_tva:.1f}%) eloigne du taux sectoriel ({btva:.0f}%)"
    details['tva'] = {'points': pts, 'max': 20, 'label': 'Coherence TVA', 'detail': lbl}

    # 3. FIN EXERCICE (15 pts)
    p7 = fe.get('p7', 0); p30 = fe.get('p30', 0)
    if   p7 > 35: pts, lbl = 15, f"{p7:.1f}% du volume sur les 7 derniers jours (fort)"
    elif p7 > 25: pts, lbl = 10, f"{p7:.1f}% du volume sur les 7 derniers jours"
    elif p30 > 50: pts, lbl =  6, f"{p30:.1f}% du volume sur les 30 derniers jours"
    else:          pts, lbl =  0, f"Repartition normale ({p7:.1f}% sur 7j)"
    details['fin_exercice'] = {'points': pts, 'max': 15, 'label': 'Concentration fin exercice', 'detail': lbl}

    # 4. ABERRANTS (10 pts)
    nb_ab = ab.get('nb', 0)
    if   nb_ab >= 8: pts, lbl = 10, f"{nb_ab} montants > 3 sigma"
    elif nb_ab >= 4: pts, lbl =  6, f"{nb_ab} montants aberrants"
    elif nb_ab >= 1: pts, lbl =  3, f"{nb_ab} montant(s) a justifier"
    else:            pts, lbl =  0, "Aucun montant aberrant"
    details['aberrants'] = {'points': pts, 'max': 10, 'label': 'Montants aberrants', 'detail': lbl}

    # 5. MARGE vs SECTEUR (12 pts)
    tm_e = bench.get('tm_entreprise'); tm_med = bench.get('tm_mediane', bench.get('tm_med', 0))
    tm_low = bench.get('tm_low', 0)
    pts = 0; lbl = "Marge dans la norme sectorielle"
    if tm_e is not None and tm_med > 0:
        ec_rel = abs(tm_e - tm_med) / tm_med * 100
        if tm_e < tm_low and ec_rel > 30:
            pts, lbl = 12, f"Marge {tm_e:.1f}% tres inferieure a mediane {tm_med:.0f}%"
        elif tm_e < tm_low:
            pts, lbl =  7, f"Marge {tm_e:.1f}% sous le bas de fourchette ({tm_low:.0f}%)"
        elif ec_rel > 40:
            pts, lbl =  5, f"Marge tres atypique (ecart {ec_rel:.0f}% vs secteur)"
    details['marge'] = {'points': pts, 'max': 12, 'label': 'Marge vs secteur', 'detail': lbl}

    # 6. IS (10 pts)
    tx_is = fisc.get('tx_is', 0); is_d = fisc.get('is_d', 0); res_f = fisc.get('res', 0)
    pts = 0; lbl = "Position IS normale"
    if res_f > 100000 and is_d == 0:
        pts, lbl = 10, f"Resultat {res_f:,.0f} EUR sans IS comptabilise"
    elif tx_is > 35:
        pts, lbl =  8, f"Taux IS apparent {tx_is:.1f}% (norme 15-25%)"
    elif tx_is < 5 and res_f > 50000:
        pts, lbl =  6, f"Taux IS tres faible {tx_is:.1f}% pour resultat {res_f:,.0f} EUR"
    details['is'] = {'points': pts, 'max': 10, 'label': 'Impot sur les Societes', 'detail': lbl}

    # 7. CHARGES SALARIALES (8 pts)
    rc = sal.get('ratio_cotis', 0)
    pts = 0; lbl = "Charges salariales normales"
    if rc > 0:
        if rc < 20:   pts, lbl = 8, f"Cotisations/salaires {rc:.1f}% (attendu 42-48%)"
        elif rc > 60: pts, lbl = 5, f"Cotisations/salaires {rc:.1f}% — anomalie"
        else:         lbl = f"Ratio cotisations/salaires {rc:.1f}%"
    details['salaires'] = {'points': pts, 'max': 8, 'label': 'Charges salariales', 'detail': lbl}

    # 8. INTEGRITE FEC (5 pts)
    nb_dup = dup.get('nb', 0)
    if nb_dup > 20: pts, lbl = 5, f"{nb_dup} doublons — importation multiple possible"
    elif nb_dup > 5: pts, lbl = 3, f"{nb_dup} doublons a verifier"
    else:            pts, lbl = 0, "Integrite FEC verifiee"
    details['integrite'] = {'points': pts, 'max': 5, 'label': 'Integrite du FEC', 'detail': lbl}

    # SCORE GLOBAL
    score_g = min(100, sum(v['points'] for v in details.values()))
    niveau = 'TRES FAIBLE'; conseil = GRILLE_RISQUE_FISCAL[-1][2]
    for seuil, lbl_n, cons in GRILLE_RISQUE_FISCAL:
        if score_g >= seuil:
            niveau = lbl_n; conseil = cons; break

    resultats = {
        'score_global':   score_g,
        'niveau_risque':  niveau,
        'conseil':        conseil,
        'details':        details,
        'nb_axes_risque': sum(1 for v in details.values() if v['points'] > 0),
        'max_theorique':  sum(v['max'] for v in details.values()),
    }
    print(f"  [MODULE 2] Bouclier Fiscal — Score : {score_g}/100 ({niveau})")
    return [], resultats


# ══════════════════════════════════════════════════════════════
# MODULE 3 — CONVERTISSEUR FEC EN PREVISIONNEL N+1
# ══════════════════════════════════════════════════════════════

def generer_previsionnel(df, stats, secteur_code,
                         croissance_ca=0.05, output_docx=None, filepath=None):
    """
    MODULE 3 — Previsionnel N+1.
    Projette CA, charges, IS sur l'annee N+1 et exporte en Word.
    """
    def s(pfx, col):
        return float(df[df['CompteNum'].str.startswith(pfx)][col].sum())

    mg  = stats.get('marges', {});    sal = stats.get('charges_salariales', {})
    fe  = stats.get('fin_exercice', {}); bench = stats.get('benchmark', {})
    sais = stats.get('saisonnalite', {})

    ca_n     = mg.get('ca',    s('70','Credit') - s('70','Debit'))
    achats_n = mg.get('achats', s('60','Debit')  - s('60','Credit'))
    cp_n     = mg.get('cp',    s('64','Debit'))
    ce_n     = mg.get('ce',    s('61','Debit') + s('62','Debit'))
    amo_n    = s('681','Debit') + s('6811','Debit')
    autres_n = s('63','Debit') + s('65','Debit') + s('67','Debit')
    res_n    = mg.get('res', 0)

    def r(v): return v / ca_n if ca_n > 0 else 0

    ca_n1    = ca_n * (1 + croissance_ca)
    achats_n1 = ca_n1 * r(achats_n)
    cp_n1    = ca_n1 * r(cp_n) * 1.03
    ce_n1    = ca_n1 * r(ce_n)
    amo_n1   = amo_n
    autres_n1 = ca_n1 * r(autres_n)
    mb_n1    = ca_n1 - achats_n1
    ebe_n1   = mb_n1 - ce_n1 - cp_n1
    rex_n1   = ebe_n1 - amo_n1 - autres_n1
    taux_is  = 0.25 if rex_n1 > 42500 else 0.15
    is_n1    = rex_n1 * taux_is if rex_n1 > 0 else 0
    res_n1   = rex_n1 - is_n1

    ca_m_n   = sais.get('ca_mensuel', {})
    ca_m_n1  = {}
    if ca_m_n and ca_n > 0:
        for mois, val in ca_m_n.items():
            ca_m_n1[mois] = round(ca_n1 * val / ca_n, 0)

    prev = {
        'annee_n':          (fe.get('date_cloture','N') or 'N')[:4],
        'secteur':          BENCHMARKS.get(secteur_code, {}).get('label', secteur_code),
        'croissance_ca_hyp': croissance_ca,
        'ca_n': ca_n, 'achats_n': achats_n, 'cp_n': cp_n,
        'ce_n': ce_n, 'dot_amo_n': amo_n,   'res_n': res_n,
        'ca_n1': ca_n1,      'achats_n1': achats_n1, 'mb_n1': mb_n1,
        'cp_n1': cp_n1,      'ce_n1': ce_n1,         'dot_amo_n1': amo_n1,
        'autres_ch_n1': autres_n1, 'ebe_n1': ebe_n1,  'rex_n1': rex_n1,
        'is_n1': is_n1,      'res_net_n1': res_n1,    'taux_is': taux_is,
        'tm_n1':   mb_n1 / ca_n1 * 100 if ca_n1 > 0 else 0,
        'tr_n1':   res_n1 / ca_n1 * 100 if ca_n1 > 0 else 0,
        'cp_ca_n1': cp_n1 / ca_n1 * 100 if ca_n1 > 0 else 0,
        'nb_sal_estime': sal.get('nb_sal_estime', 0),
        'ca_mensuel_n':  ca_m_n,
        'ca_mensuel_n1': ca_m_n1,
    }

    print(f"  [MODULE 3] Previsionnel N+1 : CA={ca_n1:,.0f} EUR "
          f"(+{croissance_ca*100:.1f}%), Res.={res_n1:,.0f} EUR")

    if output_docx is None and filepath:
        output_docx = os.path.splitext(filepath)[0] + "_previsionnel_n1.docx"

    if output_docx:
        if DOCX_AVAILABLE:
            _exporter_previsionnel_docx(prev, output_docx)
        else:
            txt = os.path.splitext(output_docx)[0] + ".txt"
            _exporter_previsionnel_txt(prev, txt)
            output_docx = txt

    return output_docx, prev


def _exporter_previsionnel_docx(prev, path):
    try:
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        an  = prev['annee_n']
        an1 = str(int(an)+1) if an.isdigit() else 'N+1'
        doc = _DocxDocument()
        for sec in doc.sections:
            sec.top_margin = sec.bottom_margin = Cm(2.5)
            sec.left_margin = Cm(3.0); sec.right_margin = Cm(2.5)

        t = doc.add_heading(f"Previsionnel {an1} — Base FEC {an}", 0)
        t.runs[0].font.color.rgb = RGBColor(0x1A,0x3A,0x5C)
        p = doc.add_paragraph()
        p.add_run(
            f"Secteur : {prev['secteur']}    "
            f"Hypothese CA : +{prev['croissance_ca_hyp']*100:.1f}%    "
            f"Genere le {datetime.now().strftime('%d/%m/%Y')}"
        ).font.size = Pt(9)
        doc.add_paragraph()

        doc.add_heading("1. Hypotheses de travail", level=1)
        for h in [
            f"Croissance CA                     : +{prev['croissance_ca_hyp']*100:.1f}%",
            "Evolution masse salariale         : +3,0%",
            f"Taux IS retenu                    : {prev['taux_is']*100:.0f}%",
            "Amortissements                    : identiques a N",
            "Charges externes                  : proportionnelles au CA",
        ]:
            doc.add_paragraph(h, style="List Bullet")
        doc.add_paragraph()

        doc.add_heading(f"2. Compte de resultat {an} / {an1}", level=1)
        tbl = doc.add_table(rows=1, cols=4); tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i,t in enumerate(["Poste", f"{an} (reel)", f"{an1} (prev.)", "Variation"]):
            hdr[i].text = t; hdr[i].paragraphs[0].runs[0].bold = True

        ca_n  = prev['ca_n']; ca_n1 = prev['ca_n1']
        lignes = [
            ("Chiffre d'affaires",       ca_n,                              ca_n1),
            ("  - Achats / CAMV",        prev['achats_n'],                  prev['achats_n1']),
            ("= Marge brute",            ca_n - prev['achats_n'],           prev['mb_n1']),
            ("  - Charges externes",     prev['ce_n'],                      prev['ce_n1']),
            ("  - Charges de personnel", prev['cp_n'],                      prev['cp_n1']),
            ("= EBE",                    ca_n-prev['achats_n']-prev['ce_n']-prev['cp_n'], prev['ebe_n1']),
            ("  - Dotations amort.",     prev['dot_amo_n'],                 prev['dot_amo_n1']),
            ("= Resultat exploitation",  prev['res_n'],                     prev['rex_n1']),
            ("  - IS",                   0,                                 prev['is_n1']),
            ("= Resultat net",           prev['res_n'],                     prev['res_net_n1']),
        ]
        for lib, vn, vn1 in lignes:
            row = tbl.add_row().cells
            row[0].text = lib; row[1].text = f"{vn:,.0f}"; row[2].text = f"{vn1:,.0f}"
            row[3].text = (f"{'+'if vn1>=vn else ''}{vn1-vn:,.0f} ({(vn1-vn)/vn*100:.1f}%)"
                           if vn != 0 else "—")
        doc.add_paragraph()

        doc.add_heading("3. Ratios cles N+1", level=1)
        for r in [
            f"Taux marge brute    : {prev['tm_n1']:.1f}%",
            f"Taux resultat net   : {prev['tr_n1']:.1f}%",
            f"Charges pers. / CA  : {prev['cp_ca_n1']:.1f}%",
            f"Effectif estime     : ~{prev['nb_sal_estime']} salarie(s)",
        ]:
            doc.add_paragraph(r, style="List Bullet")

        if prev.get('ca_mensuel_n1'):
            doc.add_paragraph()
            doc.add_heading("4. Repartition mensuelle N+1", level=1)
            MOIS = {1:'Jan',2:'Fev',3:'Mar',4:'Avr',5:'Mai',6:'Juin',
                    7:'Jul',8:'Aou',9:'Sep',10:'Oct',11:'Nov',12:'Dec'}
            t2 = doc.add_table(rows=1, cols=3); t2.style = "Table Grid"
            h2 = t2.rows[0].cells
            for i, tx in enumerate(["Mois","CA N (reel)","CA N+1 (prev.)"]): 
                h2[i].text = tx; h2[i].paragraphs[0].runs[0].bold = True
            cm_n = prev.get('ca_mensuel_n', {})
            for ms, vn1 in sorted(prev['ca_mensuel_n1'].items()):
                try: ml = MOIS.get(int(ms), ms)
                except: ml = ms
                ro = t2.add_row().cells
                ro[0].text = ml
                ro[1].text = f"{cm_n.get(ms,0):,.0f}" if cm_n else "—"
                ro[2].text = f"{vn1:,.0f}"

        doc.add_paragraph()
        pw = doc.add_paragraph(
            "AVERTISSEMENT : Previsionnel indicatif genere automatiquement. "
            "A valider par l'expert-comptable avant toute utilisation externe.")
        pw.runs[0].font.size = Pt(8); pw.runs[0].italic = True
        doc.save(path)
        print(f"  Previsionnel Word : {path}")
    except Exception as e:
        print(f"  Previsionnel docx echoue ({e})")


def _exporter_previsionnel_txt(prev, path):
    an  = prev['annee_n']
    an1 = str(int(an)+1) if an.isdigit() else 'N+1'
    lines = [
        f"PREVISIONNEL {an1} — base FEC {an}",
        f"Secteur : {prev['secteur']}  |  CA +{prev['croissance_ca_hyp']*100:.1f}%",
        "="*60,
        f"{'Poste':<32} {an:>10} {an1:>12}",
        "-"*60,
        f"{'Chiffre d affaires':<32} {prev['ca_n']:>10,.0f} {prev['ca_n1']:>12,.0f}",
        f"{'- Achats':<32} {prev['achats_n']:>10,.0f} {prev['achats_n1']:>12,.0f}",
        f"{'= Marge brute':<32} {prev['ca_n']-prev['achats_n']:>10,.0f} {prev['mb_n1']:>12,.0f}",
        f"{'- Charges externes':<32} {prev['ce_n']:>10,.0f} {prev['ce_n1']:>12,.0f}",
        f"{'- Charges de personnel':<32} {prev['cp_n']:>10,.0f} {prev['cp_n1']:>12,.0f}",
        f"{'= EBE':<32} {'':>10} {prev['ebe_n1']:>12,.0f}",
        f"{'- Amortissements':<32} {prev['dot_amo_n']:>10,.0f} {prev['dot_amo_n1']:>12,.0f}",
        f"{'= Resultat exploitation':<32} {prev['res_n']:>10,.0f} {prev['rex_n1']:>12,.0f}",
        f"{'- IS':<32} {'':>10} {prev['is_n1']:>12,.0f}",
        f"{'= Resultat net':<32} {prev['res_n']:>10,.0f} {prev['res_net_n1']:>12,.0f}",
        "="*60,
        f"Taux marge brute N+1  : {prev['tm_n1']:.1f}%",
        f"Taux resultat net N+1 : {prev['tr_n1']:.1f}%",
        f"Charges personnel/CA  : {prev['cp_ca_n1']:.1f}%",
    ]
    with open(path,'w',encoding='utf-8') as f: f.write("\n".join(lines))
    print(f"  Previsionnel TXT : {path}")


# ══════════════════════════════════════════════════════════════
# FIN DES MODULES V5
# ══════════════════════════════════════════════════════════════

def analyser_fec(filepath, output_path=None, secteur_force=None):
    print(f"\n{'='*62}\n  FEC AUDIT PRO v5.0\n  {os.path.basename(filepath)}\n{'='*62}")
    df = charger_fec(filepath)
    print(f"\n  {len(df):,} lignes chargees\n")

    # — Detection et confirmation du secteur —
    print("  Detection du secteur d'activite...")
    secteur_code, conf, _ = detecter_secteur(df)
    print(f"  Secteur estime : {BENCHMARKS[secteur_code]['label']} (confiance {conf}%)")

    if secteur_force:
        secteur_code = secteur_force
        print(f"  Secteur force  : {BENCHMARKS.get(secteur_code, {}).get('label', secteur_code)}")
    else:
        secteur_code = demander_secteur(secteur_code, conf)

    print(f"\n  Secteur retenu : {BENCHMARKS.get(secteur_code, {}).get('label', secteur_code)}\n")

    # — Controles —
    all_a = []; stats = {}
    controles = [
        ("Equilibre D/C",        analyser_equilibre,         'equilibre'),
        ("Loi de Benford",       analyser_benford,           'benford'),
        ("Montants ronds",       analyser_montants_ronds,    'montants_ronds'),
        ("Week-end",             analyser_weekend,           'weekend'),
        ("Soldes anormaux",      analyser_soldes,            'soldes'),
        ("Doublons",             analyser_doublons,          'doublons'),
        ("Ecritures inversees",  analyser_inversees,         'inversees'),
        ("Fin d exercice",       analyser_fin_exercice,      'fin_exercice'),
        ("Concentration",        analyser_concentration,     'concentration'),
        ("Saisonnalite",         analyser_saisonnalite,      'saisonnalite'),
        ("Marges et TVA",        analyser_marges,            'marges'),
        ("Montants aberrants",   analyser_aberrants,         'aberrants'),
    ]
    for nom, fn, key in controles:
        try:
            a, s = fn(df)
            all_a.extend(a); stats[key] = s
            print(f"  {'!!' if a else 'OK'} {nom:<35} {len(a)} anomalie(s)")
        except Exception as e:
            print(f"  ?? {nom:<35} Erreur: {e}")

    # — Nouveaux modules v4 —
    nouveaux = [
        ("Benchmark sectoriel",   lambda d: analyser_benchmark(d, secteur_code),   'benchmark'),
        ("Charges salariales",    analyser_charges_salariales,                     'charges_salariales'),
        ("Scoring fiscal",        analyser_scoring_fiscal,                          'scoring_fiscal'),
    ]
    for nom, fn, key in nouveaux:
        try:
            a, s = fn(df)
            all_a.extend(a); stats[key] = s
            print(f"  {'!!' if a else 'OK'} {nom:<35} {len(a)} anomalie(s)")
        except Exception as e:
            print(f"  ?? {nom:<35} Erreur: {e}")

    # — MODULE 2 v5 : Bouclier Fiscal —
    try:
        stats['_anomalies_raw'] = all_a  # reference pour le bouclier
        _, bouclier = calculer_bouclier_fiscal(df, stats, all_a, secteur_code)
        stats['bouclier_fiscal'] = bouclier
    except Exception as e:
        print(f"  ?? Bouclier Fiscal                   Erreur: {e}")

    ordre = {'CRITIQUE': 0, 'ALERTE': 1, 'ATTENTION': 2, 'INFO': 3}
    all_a.sort(key=lambda x: ordre.get(x['gravite'], 4))
    nc = sum(1 for a in all_a if a['gravite'] == 'CRITIQUE')
    print(f"\n  Total : {len(all_a)} anomalies ({nc} critiques)\n")

    print("  Generation des commentaires et benchmarks...")
    commentaires, suggestions = generer_commentaires(df, all_a, stats, secteur_code)
    html = generer_rapport(filepath, df, all_a, stats, commentaires, suggestions, secteur_code)

    if output_path is None:
        output_path = os.path.splitext(filepath)[0] + "_audit_v5.html"
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html)
    print(f"  Rapport HTML : {output_path}")

    # — MODULE 3 v5 : Previsionnel N+1 —
    no_prev = os.environ.get('FEC_NO_PREVISIONNEL', '0') == '1'
    if not no_prev:
        try:
            croissance = float(os.environ.get('FEC_CROISSANCE_CA', '0.05'))
            prev_path, prev_data = generer_previsionnel(
                df, stats, secteur_code,
                croissance_ca=croissance,
                filepath=filepath
            )
            stats['previsionnel'] = prev_data
            if prev_path:
                print(f"  Previsionnel  : {prev_path}")
        except Exception as e:
            print(f"  ?? Previsionnel N+1 : Erreur: {e}")

    # — MODULE 1 v5 : Rapport de Mission IA (appel optionnel) —
    api_key = os.environ.get('ANTHROPIC_API_KEY', '')
    if api_key:
        try:
            rapport_path, _ = generer_rapport_mission_ia(
                filepath, stats, all_a, secteur_code, api_key=api_key
            )
            if rapport_path:
                print(f"  Rapport IA    : {rapport_path}")
        except Exception as e:
            print(f"  ?? Rapport Mission IA : Erreur: {e}")
    else:
        print(f"  [MODULE 1] Rapport Mission IA : definir ANTHROPIC_API_KEY pour activer")

    print(f"\n{'='*62}\n")
    return output_path, all_a


if __name__ == "__main__":
    p = argparse.ArgumentParser(description='FEC Audit Pro v5.0')
    p.add_argument('fichier')
    p.add_argument('-o', '--output', default=None, help='Chemin du rapport HTML de sortie')
    p.add_argument('-s', '--secteur', default=None,
                   choices=[k for k in BENCHMARKS if k != 'INCONNU'] + ['INCONNU'],
                   help='Forcer le secteur (evite le dialogue interactif)')
    p.add_argument('--api-key', default=None,
                   help='Cle API Anthropic pour le rapport de mission IA (Module 1). '
                        'Peut aussi etre definie via ANTHROPIC_API_KEY.')
    p.add_argument('--croissance', type=float, default=5.0,
                   help='Hypothese de croissance du CA pour le previsionnel N+1 en %% (defaut: 5.0)')
    p.add_argument('--no-previsionnel', action='store_true',
                   help='Desactiver la generation du previsionnel N+1 (Module 3)')
    p.add_argument('--no-rapport-ia', action='store_true',
                   help='Desactiver la generation du rapport de mission IA (Module 1)')
    args = p.parse_args()

    if not os.path.exists(args.fichier):
        print(f"Fichier introuvable : {args.fichier}"); sys.exit(1)

    # Injection de la cle API si fournie en argument
    if args.api_key:
        os.environ['ANTHROPIC_API_KEY'] = args.api_key
    if args.no_rapport_ia:
        os.environ.pop('ANTHROPIC_API_KEY', None)

    # Stockage de l'hypothese de croissance dans l'environnement pour analyser_fec
    os.environ['FEC_CROISSANCE_CA'] = str(args.croissance / 100)
    os.environ['FEC_NO_PREVISIONNEL'] = '1' if args.no_previsionnel else '0'

    analyser_fec(args.fichier, args.output, secteur_force=args.secteur)
